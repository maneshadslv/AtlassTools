import os
import pandas as pd
import shapefile_old as shp
from shapely.geometry import Point
from shapely.geometry import Polygon
from shapely.ops import unary_union
import statistics
import subprocess
import numpy as np
import time
import docxtpl
from docx.shared import Pt
from docx.shared import Mm
import jinja2
from datetime import datetime
import laspy
import matplotlib.pyplot as plt
import seaborn as sns
from gooey import GooeyParser, Gooey
import struct

density_report_template = r"Z:\Reporting_Templates\Density\Density_Reporting_Template.docx"

test_pt = r"Y:\Coverage_Setup\BR02284_Darwin\gronly\buffered_tiles\test_pt.shp"


class coverage_reporter:
    def __init__(self, sa_folder, aoi, hydro_poly, gm, user, br, cores, area_name, tile_s, req_density, method_in,
                 data_source, run_final_only, zone):
        self.sa_folder = sa_folder
        self.aoi = aoi
        print('aoi is', self.aoi)
        if self.aoi:
            self.imported_aoi = shp.Reader(self.aoi)
            print(self.imported_aoi)
        else:
            self.imported_aoi = None
        self.hydro_poly = hydro_poly
        self.gm_exe = gm
        self.user = user
        self.br_code = br
        self.cores = cores
        self.area_name = area_name
        self.tile_size = tile_s
        self.req_d = float(req_density)
        self.method = method_in
        self.run_final_only = run_final_only
        self.cell_size = 5
        self.lower_limit = 0  # catches stray values; needs to be monitored
        self.upper_limit = 200
        self.concavity = 5
        self.compliant_cells = None
        self.hydro_area = 0
        if 'yes' in data_source or 'Yes' in data_source:
            self.data_already_tiled = True
        else:
            self.data_already_tiled = False
        self.zone = zone
        self.donut_hydro = False

        self.density_working_folder = os.path.join(self.sa_folder, 'density_working')
        self.hydro_working_folder = os.path.join(self.sa_folder, 'hydro_clipping')
        self.density_laz_folder = os.path.join(self.density_working_folder, 'density_laz')
        if not os.path.isdir(self.density_working_folder):
            os.mkdir(self.density_working_folder)
        if not os.path.isdir(self.density_laz_folder):
            os.mkdir(self.density_laz_folder)
        if not os.path.isdir(self.hydro_working_folder) and self.hydro_poly is not None:
            os.mkdir(self.hydro_working_folder)

        if self.hydro_poly is not None:
            self.imported_hydro = shp.Reader(self.hydro_poly)
            self.hydro_as_poly = None
            self.import_hydro()

        if not self.run_final_only:
            if not self.data_already_tiled:
                # step 1: lasindex swaths
                self.index_laz()
                # step 2: tile swaths
                self.tile_laz()
                self.tiled_laz_folder = self.density_working_folder
            else:
                self.tiled_laz_folder = self.sa_folder

            # step 3: lasgrid
            self.make_grid()

            # step 3: hydro
            if self.hydro_poly is not None:

                # step 4: class hydro points out
                self.lasclip_hydro()

            else:
                pass

        if self.hydro_poly is not None:
            self.density_laz_folder = self.hydro_working_folder

        # step 4: import AOI
        self.aoi_as_poly = None
        if self.aoi:
            self.import_aoi()

        # step 5: lasclip grid results
        if self.aoi and not self.run_final_only:
            self.run_lasclip(self.density_laz_folder)
        elif not self.run_final_only:
            self.l2l_fakeout()
        # step 6: laspy > pandas
        self.density_laz_list = [os.path.join(self.density_laz_folder, a) for a in os.listdir(self.density_laz_folder)
                                 if a.endswith('0_clipped.laz')]
        self.unfiltered_density_pandas = pd.DataFrame()
        for i, laz in enumerate(self.density_laz_list):

            this_pd = self.laspy_to_pandas(laz, i)
            if self.unfiltered_density_pandas.empty:
                self.unfiltered_density_pandas = this_pd
            else:
                t = self.unfiltered_density_pandas
                self.unfiltered_density_pandas = pd.concat([t, this_pd], axis=0, ignore_index=True)

        # vestigial, fix later
        self.filtered_density_pandas = self.unfiltered_density_pandas

        # coverage checking
        self.cov_area = len(self.filtered_density_pandas) * (self.cell_size * self.cell_size)
        if self.aoi:
            self.aoi_area = self.aoi_as_poly.area
            if self.aoi_area == 0:
                print('Error with AOI, area issue')
                print(self.aoi_as_poly)
                exit()
            self.area_difference = self.aoi_area - self.cov_area
        else:
            self.aoi_area = None
            self.area_difference = None

        # step 9: analyze results; add compliance column to pandas
        self.information_table = pd.DataFrame()
        self.st_dev_density = None
        self.analyse_pandas()
        self.histo_out = os.path.join(self.density_laz_folder, 'density_histo.png')
        self.make_density_histo()

        # step 10: make image for report
        self.image_path = os.path.join(self.density_laz_folder, 'output_density_image.jpg')
        self.make_report_image()

        # step 11: make report
        self.report_path = os.path.join(self.density_working_folder, '%s_%s_Density_Report.docx' % (self.br_code, self.area_name))
        self.make_report()

    def index_laz(self):
        print('Indexing laz for %s...' % self.area_name)
        command_string = ['lasindex', '-i', str(self.sa_folder) + r'\*.laz', '-cores', str(self.cores)]
        subprocess.call(command_string)

    def l2l_fakeout(self):
        command_string = ['las2las', '-i', str(self.density_laz_folder) + r'\*0.laz', '-cores', str(self.cores), '-odix',
                          '_clipped', '-olaz']
        subprocess.call(command_string)

    def tile_laz(self):
        print('Tiling laz for %s...' % self.area_name)
        command_string = ['lastile', '-i', str(self.sa_folder) + r'\*.laz', '-merged', '-tile_size', str(self.tile_size), '-olaz',
                          '-odir', str(self.density_working_folder)]
        subprocess.call(command_string)

    def make_grid(self):
        print('Building last echo grid for %s...' % self.area_name)
        command_string = ['lasgrid', '-i', str(self.tiled_laz_folder) + r'\*.laz', '-odir', self.density_working_folder,
                          '-odir', str(self.density_laz_folder), '-olaz', '-step', str(self.cell_size), '-point_density_32bit',
                          '-cores', str(self.cores), '-last_only']
        subprocess.call(command_string)

    def import_aoi(self):
        print('Importing AOI...')
        poly_list = []
        print(self.imported_aoi)
        try:
            for s in self.imported_aoi.shapeRecords():
                coords = s.__geo_interface__['geometry']['coordinates']
                if len(coords) == 1 and type(coords[0]) is tuple:
                    coords = coords[0]
                    try:
                        s_p = Polygon(coords)
                        poly_list.append(s_p)
                    except AssertionError:
                        for p in coords:
                            s_p = Polygon(p)
                            poly_list.append(s_p)
                else:
                    for g in coords:

                        if len(g) == 1:
                            print('len g 1')
                            s_p = Polygon(g[0])
                            poly_list.append(s_p)
                        else:
                            for k in g:
                                print('k found, len ', len(k))
                                s_p = Polygon(k[0])
                                poly_list.append(s_p)
        except struct.error:
            print(self.imported_aoi.shapeRecords()[0].__geointerface__)
            exit()

        self.aoi_as_poly = unary_union(poly_list)

    def import_hydro(self):
        print('Importing hydro...')
        poly_list = []
        outer = None
        for s in self.imported_hydro.shapeRecords():
            coords = s.__geo_interface__['geometry']['coordinates']
            print(s.__geo_interface__['geometry']['type'])
            if len(coords) == 1 and type(coords[0]) is tuple:
                coords = coords[0]
                try:
                    s_p = Polygon(coords)
                    poly_list.append(s_p)
                except AssertionError:
                    for p in coords:
                        s_p = Polygon(p)
                        poly_list.append(s_p)
            else:
                outer = Polygon(coords[0])
                for i, g in enumerate(coords[1:]):
                    s_p = Polygon(g)
                    poly_list.append(s_p)
        if outer is None:
            self.hydro_as_poly = unary_union(poly_list)
        else:
            subtraction = unary_union(poly_list)
            self.hydro_as_poly = outer.difference(subtraction)
            self.donut_hydro = True
        self.hydro_area = self.hydro_as_poly.area

    def lasclip_hydro(self):
        print('Clipping hydro points...')
        command_string = ['lasclip', '-i', str(self.density_laz_folder) + r'\*0.laz', '-cores', str(self.cores), '-poly',
                          str(self.hydro_poly), '-olaz', '-odir', str(self.hydro_working_folder),
                          '-classify', '25', '-odix', '_hydro']
        if self.donut_hydro:
            command_string.append('-donuts')
        subprocess.call(command_string)
        # donut doesn't seem to work grr
        print('Prepping for AOI clip...')
        command_string_2 = ['las2las', '-i', str(self.hydro_working_folder) + r'\*_hydro.laz', '-keep_class', '25', '-olaz', '-cores', str(self.cores),
                            '-odix', '_sifted0']
        subprocess.call(command_string_2)

    def exclude_hydro(self):
        temporary_unfiltered = self.unfiltered_density_pandas
        for i in temporary_unfiltered.index.values:
            print('Checking out %s of %s values...' % (i, len(temporary_unfiltered)))
            x = temporary_unfiltered.iloc[i]['Easting']
            y = temporary_unfiltered.iloc[i]['Northing']
            i_as_pt = Point([x, y])
            if i_as_pt.within(self.hydro_as_poly):
                temporary_unfiltered.drop(temporary_unfiltered.index[i])
        self.filtered_density_pandas = temporary_unfiltered

    def import_cov_shp(self, shape_file_path):
        s_in = shp.Reader(shape_file_path).shapes()
        if len(s_in) > 0:
            poly_list = []
            for s in s_in:
                coords = s.__geo_interface__['coordinates']
                if len(coords) == 1 and type(coords[0]) is tuple:
                    coords = coords[0]
                s_p = Polygon(coords)
                poly_list.append(s_p)
            s_in_as_poly_rgh = unary_union(poly_list)
        else:
            s_in_as_poly_rgh = None
        return s_in_as_poly_rgh

    def run_lasclip(self, folder_in):
        print('Clipping density laz to AOI...')
        command_string = ['lasclip', '-i', str(folder_in) + r'\*0.laz', '-cores', str(self.cores), '-odix',
                          '_clipped', '-poly', str(self.aoi), '-olaz']
        subprocess.call(command_string)

    def run_lasbound(self):
        print('Running lasboundary to find coverage...')
        command_string = ['lasboundary', '-i', str(self.density_working_folder) + r'\*0_clipped.laz', '-cores',
                          str(self.cores), '-oshp', '-concavity', str(self.concavity), '-holes']
        subprocess.call(command_string)

    '''def run_lasbound_on_all(self):
        print('Running lasboundary to find coverage on all merged...')
        working_laz = str(self.density_working_folder) + r'\*0.laz'
        print(working_laz)
        command_string = ['lasboundary', '-i', str(working_laz), '-merged', '-o',
                          str(self.cov_polygon), '-concavity', str(self.concavity), '-holes', '-disjoint']
        subprocess.call(command_string)'''

    def laspy_to_pandas(self, laz_path, count):
        print('Importing %s, file %s of %s...' % (os.path.basename(laz_path), count, len(self.density_laz_list)))
        infile = laspy.file.File(laz_path)
        headers = ['Easting', 'Northing', 'PointCount']
        coords = np.vstack((infile.x, infile.y, infile.z)).transpose().tolist()
        temp_dframe = pd.DataFrame(columns=headers, data=coords)
        temp_dframe['Density'] = temp_dframe['PointCount']  # / self.cell_size
        temp_dframe_v2 = temp_dframe[(temp_dframe['Density'] <= self.upper_limit) &
                                     (temp_dframe['Density'] >= self.lower_limit)]
        if not len(temp_dframe_v2) == len(temp_dframe):
            print('Removed stray value from %s...' % os.path.basename(laz_path))
        # print(temp_dframe_v2[['Density', 'PointCount']])
        # exit()
        return temp_dframe_v2

    def write_a_polygon(self, geometry_in):
        print('Writing a polygon...')
        out_path = os.path.join(self.density_working_folder, '%s_%s_Coverage_poly.shp' % (self.br_code, self.area_name))
        m = shp.Writer(out_path, shapeType=5)
        m.field('Item', 'N')

        for i, s in enumerate(geometry_in.__geo_interface__['coordinates']):
            coord_tuple = s[0]
            coord_list = [list(a) for a in coord_tuple]
            try:
                m.poly([coord_list])
            except Exception as e:
                print('shp except', e)
                m.null()
            m.record(Item=i)
        m.close()

    def analyse_pandas(self):
        print('Analysing results...')
        # make compliance column
        self.filtered_density_pandas['Compliance'] = np.where(self.filtered_density_pandas['Density'] >= self.req_d,
                                                              True, False)
        # print(self.filtered_density_pandas[['Density', 'Compliance', 'PointCount']])
        compliance_count = self.filtered_density_pandas['Compliance'].values.sum()
        non_compliance_count = (~self.filtered_density_pandas['Compliance']).values.sum()
        # print(compliance_count, non_compliance_count)
        non_compliant_panda = self.filtered_density_pandas[~self.filtered_density_pandas['Compliance']]
        non_compliant_areas_path = os.path.join(self.density_working_folder, "noncompliant_areas.csv")
        non_compliant_panda.to_csv(non_compliant_areas_path)
        # exit()
        compliance_pc = compliance_count / (compliance_count + non_compliance_count) * 100
        min_density = self.filtered_density_pandas['Density'].min()
        max_density = self.filtered_density_pandas['Density'].max()
        mean_density = self.filtered_density_pandas['Density'].mean()
        self.compliant_cells = compliance_count

        self.st_dev_density = self.filtered_density_pandas['Density'].std()

        pre_hydro_count = len(self.unfiltered_density_pandas)
        if self.aoi_area is not None:
            area_aoi = str(round((self.aoi_area / 1000000), 2)) + 'km2'
            pc_cov = round((((self.cov_area + self.hydro_area) / self.aoi_area) * 100), 2)
            if pc_cov > 100:
                pc_cov = 100
        else:
            area_aoi = "No AOI"
            pc_cov = "No AOI"
        print('pc covered preview:', pc_cov)
        print('aoi', self.aoi_area)
        print('hydro area', self.hydro_area)
        print('cov area', self.cov_area)
        if self.aoi_area is not None:
            if (self.cov_area + self.hydro_area) >= self.aoi_area:
                c_area = area_aoi
            else:
                c_area = str(round(((self.cov_area + self.hydro_area) / 1000000), 2)) + 'km2'
        else:
            c_area = str(round(((self.cov_area + self.hydro_area) / 1000000), 2)) + 'km2'

        data = {'Density compliance, percent': round(compliance_pc, 2), 'Minimum density': min_density,
                'Maximum density': max_density, 'Mean density': round(mean_density, 2),
                'Compliance requirement': self.req_d,
                'Standard deviation': round(self.st_dev_density, 2),
                'Area of AOI': area_aoi,
                'Covered area': c_area,
                'Percent covered': pc_cov}
        series_in = pd.Series(data, name='Results')
        ordered_headers = ['Density compliance, percent', 'Minimum density', 'Maximum density', 'Mean density',
                           'Standard deviation', 'Compliance requirement',
                           'Area of AOI', 'Covered area', 'Percent covered']
        self.information_table = pd.DataFrame(series_in).reindex(ordered_headers)
        self.information_table.reset_index(inplace=True)

    def make_density_histo(self):
        print('Making histogram...')
        mean_density = self.filtered_density_pandas['Density'].mean()
        min_val = mean_density - (3 * self.st_dev_density)
        max_val = mean_density + (3 * self.st_dev_density)

        all_density_list = self.filtered_density_pandas['Density'].to_list()
        # filtered to within 3 stdevs, as otherwise plot is a bit meh
        filtered_densities = [a for a in all_density_list if min_val < a < max_val]
        filtered_density_series = pd.Series(data=filtered_densities, name="Points per sq m")

        sns.set(style="darkgrid")
        sns.set_color_codes()
        # plt.axes(xbound=(0, 100))
        sns.distplot(filtered_density_series.dropna(), norm_hist=True, color="r")
        plt.savefig(self.histo_out)
        sns.reset_defaults()
        sns.reset_orig()
        plt.clf()

    def make_report_image(self):
        print('Making image for report...')
        gms_path = os.path.join(self.density_laz_folder, 'image_make.gms')
        gmc_path = os.path.join(self.density_laz_folder, 'image_make.gmc')
        if os.path.isfile(gmc_path):
            os.remove(gmc_path)
        shader_name = (str(self.req_d) + "PPM").replace(".", "-")

        gms_string = 'GLOBAL_MAPPER_SCRIPT VERSION=1.00 ENABLE_PROGRESS=YES LOG_TO_COMMAND_PROMPT=YES\n' \
                     'UNLOAD_ALL\n' \
                     'DEFINE_SHADER SHADER_NAME="%s" BLEND_COLORS=NO ' \
                     'STRETCH_TO_RANGE=NO SHADE_SLOPES=NO\n' \
                     '0.0, RGB(0,0,0)\n' \
                     '1.0, RGB(255,0,0)\n' \
                     '%s, RGB(12, 250, 242)\n' \
                     '%s, RGB(13,255,17)\n' \
                     'END_DEFINE_SHADER\n' \
                     'DEFINE_PROJ PROJ_NAME="MGA_ZONE%s_GDA_2020_AUSTRALIAN_GEODETIC_2020"\n' \
                     'Projection     MGA (Map Grid of Australia)\n' \
                     'Datum          GDA2020\n' \
                     'Zunits         NO\n' \
                     'Units          METERS\n' \
                     'Zone           %s\n' \
                     'Xshift         0.000000\n' \
                     'Yshift         0.000000\n' \
                     'Parameters\n' \
                     'END_DEFINE_PROJ\n' \
                     'SET_OPT LIDAR_FILTER="ALL" LIDAR_RETURN_FILTER="ALL" LIDAR_DRAW_MODE="ELEV"\n' \
                     'EDIT_MAP_CATALOG FILENAME=%s CREATE_IF_EMPTY=YES ADD_FILE=%s/*d.laz ZOOM_DISPLAY="PERCENT,0.90,0" ' \
                     'PROJ="MGA_ZONE%s_GDA_2020_AUSTRALIAN_GEODETIC_2020"\n' \
                     'SET_VERT_DISP_OPTS SHADER_NAME=%s ENABLE_HILL_SHADING=NO\n' \
                     'IMPORT FILENAME=%s TYPE=SHAPEFILE PROJ="MGA_ZONE%s_GDA_2020_AUSTRALIAN_GEODETIC_2020" ' \
                     'LAYER_DESC=AOI\n' \
                     'IMPORT FILENAME=%s \n' \
                     'EDIT_VECTOR FILENAME=AOI STYLE_ATTR="LINE_COLOR=RGB(0,0,0)"\n ' \
                     'EDIT_VECTOR FILENAME=AOI STYLE_ATTR="LINE_WIDTH=5"\n' % (shader_name, self.req_d,
                                                                               self.req_d+1, shader_name, shader_name,
                                                                               gmc_path, self.density_laz_folder,
                                                                               self.zone, shader_name,
                                                                               self.aoi, self.zone, gmc_path)
        if self.hydro_poly:
            gms_string = gms_string + 'IMPORT FILENAME=%s TYPE=SHAPEFILE USE_DEFAULT_PROJ=YES LAYER_DESC=HYDRO\n' \
                                      'EDIT_VECTOR FILENAME=HYDRO STYLE_ATTR="LINE_COLOR=RGB(0,123,255)" ' \
                                      'STYLE_ATTR="LINE_WIDTH=5" ATTR_TO_DELETE="NAME"\n' % self.hydro_poly

        gms_string = gms_string + 'EXPORT_RASTER FILENAME=%s TYPE=JPEG FILL_GAPS=YES ' \
                                  'INC_VECTOR_DATA=YES' % self.image_path
        with open(gms_path, 'w') as w:
            w.write(gms_string)
        command_string = [str(self.gm_exe), str(gms_path)]
        subprocess.call(command_string)

    def make_table_doco(self, table_pd, title):
        doc = docxtpl.DocxTemplate(density_report_template)
        styles = doc.styles
        doco = doc.new_subdoc()
        doco_out = None
        try:
            doco.add_paragraph("")
            doco.add_heading(title, level=2)
            tin = doco.add_table(table_pd.shape[0] + 1, table_pd.shape[1])
            tin.allow_autofit = True
            for j in range(1, table_pd.shape[-1], 1):
                tin.cell(0, j).text = str(table_pd.columns[j])
            for i in range(table_pd.shape[0]):
                for j in range(table_pd.shape[-1]):
                    tin.cell(i + 1, j).text = str(table_pd.values[i, j])
            for row in tin.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(9)
            doco_out = doco
        except KeyError:
            print('Could not gen table', title)
        return doco_out

    def make_report(self):
        print('Building report...')
        doc = docxtpl.DocxTemplate(density_report_template)
        styles = doc.styles
        date_now = datetime.today().strftime('%c')
        basic_info_table = self.make_table_doco(self.information_table.round(3), 'Summary')

        histo_inline = docxtpl.InlineImage(doc, self.histo_out)
        density_im_inline = docxtpl.InlineImage(doc, self.image_path, height=Mm(180))

        context = {'br_code': self.br_code, 'area_name': self.area_name, 'date_gen': date_now,
                   'cov_w_aoi': density_im_inline, 'histo': histo_inline, 'dens_table': basic_info_table,
                   'user': self.user, 'method': self.method, 'cell_count': self.compliant_cells,
                   'cell_sq': self.cell_size * self.cell_size}
        doc.render(context)
        doc.save(self.report_path)
        print('Wrote %s...' % os.path.basename(self.report_path))


class noot:
    def __init__(self, gooey_input):
        # testing_folder, testing_aoi, hydro_poly, gm_exe, user_name, br, cores, area_name,
        # tile_size, required_density_ppm, method
        self.testing_folder = gooey_input.sa_folder
        self.testing_aoi = gooey_input.aoi
        self.hydro_poly = gooey_input.hydro_poly
        '''if not os.path.isfile(self.hydro_poly):
            self.hydro_poly = None'''
        self.gm_exe = gooey_input.gm
        self.user_name = gooey_input.user
        self.br_code = gooey_input.br
        self.area_name = gooey_input.area_name
        self.cores = gooey_input.cores
        self.tile_size = gooey_input.tile_size
        self.req_d = gooey_input.req_d
        self.meth = gooey_input.meth
        self.setup = gooey_input.data_source
        self.run_final_only = gooey_input.run_final_only
        self.zone = gooey_input.zone

        coverage_reporter(self.testing_folder, self.testing_aoi, self.hydro_poly, self.gm_exe,
                          self.user_name, self.br_code, self.cores, self.area_name, self.tile_size,
                          self.req_d, self.meth, self.setup, self.run_final_only, self.zone)


@Gooey(program_name="Density Report Prepper", use_legacy_titles=True, required_cols=1, default_size=(750, 800))
def goo():
    parser = GooeyParser(description="Density Report Tool")
    parser.add_argument('-sa_folder', metavar='LAZ folder', widget='DirChooser')
    parser.add_argument('-aoi', metavar='AOI shapefile for area', widget='FileChooser')
    parser.add_argument('--hydro_poly', metavar='Hydro shp', widget='FileChooser')
    parser.add_argument('-gm', metavar='Global Mapper exe location', widget='FileChooser',
                        default=r"C:\Program Files\GlobalMapper21.0_64bit\global_mapper.exe")
    parser.add_argument('-user', metavar='Your name')
    parser.add_argument('-br', metavar='br_code')
    parser.add_argument('-cores', metavar='Cores to use')
    parser.add_argument('-area_name', metavar="Name of area")
    parser.add_argument('-tile_size', metavar='Size to tile to (m)')
    parser.add_argument('-req_d', metavar='Density specified in quote')
    parser.add_argument('-meth', metavar='Method of density calculation', choices=['per-tile', 'per-strip'],
                        default='per-tile')
    parser.add_argument('-data_source', metavar='Is your data already tiled?', choices=['Yes', 'No'])
    parser.add_argument('-zone', metavar='AOI Zone', choices=['48', '49', '50', '51', '52', '53', '54', '55', '56', '57'],
                        default='56')
    parser.add_argument('-run_final_only', metavar='Run final report only?', action='store_true')
    return parser.parse_args()


noot(goo())
