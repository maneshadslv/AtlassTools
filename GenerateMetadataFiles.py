#-----------------------------------------------------------------------------------------------------------------
#Include libraries
#-----------------------------------------------------------------------------------------------------------------
import itertools
import time
import random
import sys, getopt
import math
import shutil
import subprocess
import urllib
import os, glob
import numpy as np
from gooey import Gooey, GooeyParser
from multiprocessing import Pool,freeze_support
sys.path.append('{0}/lib/atlass/'.format(sys.path[0]).replace('\\','/'))
from Atlass_beta1 import *
import xml.etree.ElementTree as ET
from docx import Document

#-----------------------------------------------------------------------------------------------------------------
#Gooey input
#-----------------------------------------------------------------------------------------------------------------

@Gooey(program_name="Making Metadata file", advanced=True, default_size=(1000,800), use_legacy_titles=True, required_cols=1, optional_cols=3)
def param_parser():
    main_parser=GooeyParser(description="Reading XML")
    main_parser.add_argument("variable_file", metavar="Variable text File", widget="FileChooser", help="Select .txt file", default="")
    main_parser.add_argument("xml", metavar="XML template File", widget="FileChooser", help="Select .xml file", default="")
    main_parser.add_argument("wd", metavar="Word template File", widget="FileChooser", help="Select .docx file", default="")
    main_parser.add_argument("outputpath", metavar="Output Directory",widget="DirChooser", help="Output directory", default='')
    main_parser.add_argument("fid1", metavar="File ID 1", default='')
    main_parser.add_argument("fid2", metavar="File ID 2", default='')
    main_parser.add_argument("fid3", metavar="File ID 3", default='')
    main_parser.add_argument("fid4", metavar="File ID 4", default='')
    main_parser.add_argument("flownyear", metavar="Year Flown", default='2018')

    return main_parser.parse_args()

#-----------------------------------------------------------------------------------------------------------------
#Function definitions
#-----------------------------------------------------------------------------------------------------------------


#-----------------------------------------------------------------------------------------------------------------
#Main entry point
#-----------------------------------------------------------------------------------------------------------------
def main():

    print("Starting Program \n\n")
    freeze_support() 
    
    #Set Arguments
    args = param_parser()

    template_XML = args.xml
    template_WD = args.wd
    variable_file = args.variable_file
    outputdir = args.outputpath
    flownyear = args.flownyear
    fid1 = args.fid1
    fid2 = args.fid2
    fid3 = args.fid3
    fid4 = args.fid4

    myvars = {}
    with open(variable_file) as varfile:
        for line in varfile:
            var, val = line.partition(":")[::2]
            myvars[var.strip()] = val.strip()

    products = {}

    products['dem_asc'] = {}
    products['dem_asc']['uid'] = fid1
    products['dem_asc']['surface_type'] = 'DEM'
    products['dem_asc']['product_type'] = 'GRID'
    products['dem_asc']['format'] = 'ASCII'
    products['dem_asc']['xml'] = os.path.join(outputdir, "{0}_{2}_z{1}_DEM_GRID_1_ASCII.xml".format(myvars['#areaname#'],myvars['#zone#'],flownyear))
    products['dem_asc']['limitations'] = 'DEM accuracy will be limited by the spatial accuracy of the LiDAR point data and will contain some additional error due to interpolation, particularly in areas of dense vegetation where ground points are sparse. There may also be some minor error due to ground point misclassification.'

    products['dem_xyz'] = {}
    products['dem_xyz']['uid'] = fid2
    products['dem_xyz']['surface_type'] = 'DEM'
    products['dem_xyz']['product_type'] = 'GRID'
    products['dem_xyz']['format'] = 'TEXT'
    products['dem_xyz']['xml'] = os.path.join(outputdir, "{0}_{2}_z{1}_DEM_GRID_1_TEXT.xml".format(myvars['#areaname#'],myvars['#zone#'],flownyear))
    products['dem_xyz']['limitations']= 'DEM accuracy will be limited by the spatial accuracy of the LiDAR point data and will contain some additional error due to interpolation, particularly in areas of dense vegetation where ground points are sparse. There may also be some minor error due to ground point misclassification.'

    products['int'] = {}
    products['int']['uid'] = fid3
    products['int']['surface_type'] = 'INT-First'
    products['int']['product_type'] = 'Other'
    products['int']['format'] = 'TIFF'
    products['int']['xml'] = os.path.join(outputdir, "{0}_{2}_z{1}_INT-First_Other_1_TIFF.xml".format(myvars['#areaname#'],myvars['#zone#'],flownyear))
    products['int']['limitations'] = 'The intensity image accuracy will be limited by the spatial accuracy of the LiDAR point data.'

    products['las_ahd'] = {}
    products['las_ahd']['uid'] = fid4
    products['las_ahd']['surface_type'] = 'LiDAR-AHD'
    products['las_ahd']['product_type'] = 'MassPoints'
    products['las_ahd']['format'] = 'LAS'
    products['las_ahd']['xml'] = os.path.join(outputdir, "{0}_{2}_z{1}_LiDAR-AHD_MassPoints_1_LAS.xml".format(myvars['#areaname#'],myvars['#zone#'],flownyear))
    products['las_ahd']['limitations'] = 'The workflow and quality assurance processes were designed to achieve the Level 2 requirement for removal of significant anomalies which remain in the ground class (2), vegetation classes (3, 4, 5), buildings and structures (6), water (9), and bridges (10), and achieve a ground point misclassification rate of 2% or less. The classification accuracy was not measured.'

    print(myvars['#areaname#'])

    word_Doc = os.path.join(outputdir, "TMR_Metadata_{0}.docx".format(myvars['#areaname#']))

    document = Document(template_WD)    
    tables = document.tables

    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in myvars.items():
                        if key in para.text:
                            para.text = para.text.replace(key,val)


    document.save(word_Doc)

    with open(template_XML,encoding='latin-1') as myasciif:
        data = myasciif.read()
        for key, val in myvars.items():
            if key in data:
                print(key,val)
                data = data.replace(key,val)

    for product,params in products.items():
        testdata = data
        print(product)
        xmlfile = products[product]['xml']
        limitations = products[product]['limitations']
        uid = products[product]['uid']

        testdata = testdata.replace('#uid#',uid)
        testdata = testdata.replace('#limitations#',limitations)

        for key1,value1 in params.items():
            testdata = testdata.replace('#{0}#'.format(key1),value1)

        print(xmlfile)
  
        
        with open(xmlfile, 'wb') as f:
            #print(testdata)
            testdata = testdata.encode(encoding='latin-1',errors='strict')
            f.write(testdata)
            
            xmlfile = ''
    return


if __name__ == "__main__":
    main()         

