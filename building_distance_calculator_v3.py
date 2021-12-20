from shapely.geometry import MultiPoint
from shapely.geometry import MultiLineString
from shapely.geometry import LineString
from shapely.geometry import Polygon
from shapely.geometry import Point
from shapely.ops import unary_union
from shapely.ops import nearest_points
import pandas as pd
import os
import sys
sys.path.append('{0}/lib/shapefile_original/'.format(sys.path[0]).replace('\\','/'))
import shapefile_original as shp
import matplotlib.pyplot as plt
import seaborn as sns
import statistics
from scipy.stats import sem, t, ttest_rel
# from scipy import mean
import numpy as np
import docxtpl
from docx.shared import Pt
from docx.shared import Mm
from datetime import datetime
from gooey import Gooey, GooeyParser
from collections import OrderedDict

# templates
front_page = r"Z:\Reporting_Templates\Horizontal\horizontal_justification_front_page.docx"
sub_page = r"Z:\Reporting_Templates\Horizontal\horizontal_justification_subpage.docx"
psar_page = r"Z:\Reporting_Templates\Horizontal\TEMPLATE_PSAR_HorizontalControl_Section.docx"
divergence_dic = {'VQ780i': 0.18, 'H68': 0.60, 'VQ780ii': 0.18}


class tool:
    def __init__(self, noot):
        self.analysis_folder = noot.af
        self.flying_height = float(noot.fh)
        self.system = noot.st
        self.br = noot.br
        self.area_name = noot.an

        self.divergence = divergence_dic[self.system]
        self.ft = (self.divergence * self.flying_height)/1000
        self.poly_list = [os.path.join(self.analysis_folder, a) for a in os.listdir(self.analysis_folder) if
                          a.endswith('.shp')
                          and 'POLY_' in a]
        self.points_list = [os.path.join(self.analysis_folder, a) for a in os.listdir(self.analysis_folder) if
                            a.endswith('.shp')
                            and 'POINTS_' in a]
        self.image_list = [os.path.join(self.analysis_folder, a) for a in os.listdir(self.analysis_folder) if
                           a.endswith('.jpg')]

        oh_no = False
        if len(self.poly_list) == 0:
            print("I can't find any buildings here. Perhaps look elsewhere?")
            oh_no = True
        if len(self.points_list) == 0:
            print("I can't see any LiDAR files here. Have we checked locations and naming?")
            oh_no = True
        if oh_no:
            exit()

        self.reprots_folder = os.path.join(self.analysis_folder, 'Reports')
        if not os.path.isdir(self.reprots_folder):
            os.mkdir(self.reprots_folder)

        self.poly_shp_dic = {}
        self.poly_path_to_points_dic = {}
        self.name_to_poly = {}
        self.poly_to_points_list = {}
        self.poly_to_points_list_POST = {}
        self.useful_buildings = []

        self.points_pd = pd.DataFrame()

        # for each poly file:
        #  - load it
        #  - load corresponding pts file
        #  - find pairs
        #  - find centroids
        #  - save per pt info to pd
        #  - save centroid info to dic
        headers = ['Building Name', 'Building ID', 'Pairs', 'Original LiDAR Centroid',
                   'Original Building Centroid', 'Original Offset X', 'Original Offset Y']
        per_building_summary_list = []
        for i, poly_file in enumerate(self.poly_list):
            self.load_poly_file(poly_file)
            building_name = poly_file.split('__')[-1].replace('_shp.shp', '')
            self.name_to_poly[building_name] = poly_file
            if poly_file in self.poly_path_to_points_dic.keys():
                # get info
                points_file = self.poly_path_to_points_dic[poly_file]
                points_list = self.load_lidar_points(points_file)
                self.poly_to_points_list[poly_file] = points_list
                # find pairs, compile
                building_pd = self.find_pairs(building_name, 'prelim')
                building_pd.loc[:, 'Building ID'] = i
                if self.points_pd.empty:
                    self.points_pd = building_pd
                else:
                    temp = self.points_pd
                    self.points_pd = pd.concat([temp, building_pd], ignore_index=True, sort=False)
                # find centroids
                lidar_points_as_MP = MultiPoint(points_list)
                poly_points_as_MP = MultiPoint(building_pd['Nearest Poly Pt'].to_list())
                lidar_points_centroid = lidar_points_as_MP.centroid
                lidar_points_centroid_x = lidar_points_centroid.x
                lidar_points_centroid_y = lidar_points_centroid.y
                poly_points_centroid = poly_points_as_MP.centroid
                poly_points_centroid_x = poly_points_centroid.x
                poly_points_centroid_y = poly_points_centroid.y
                centroid_offset_x = poly_points_centroid_x - lidar_points_centroid_x
                centroid_offset_y = poly_points_centroid_y - lidar_points_centroid_y
                summary_info = [building_name, i, len(points_list), (lidar_points_centroid_x, lidar_points_centroid_y),
                                (poly_points_centroid_x, poly_points_centroid_y), centroid_offset_x, centroid_offset_y]
                per_building_summary_list.append(summary_info)
        self.initial_summary_info = pd.DataFrame(columns=headers, data=per_building_summary_list)
        # find mean offset
        print('Finding initial centroid offsets...')
        self.mean_x_centroid_offset = self.initial_summary_info['Original Offset X'].mean()
        self.mean_y_centroid_offset = self.initial_summary_info['Original Offset Y'].mean()
        self.initial_x_rmse, self.initial_x_ci95, self.initial_y_rmse, self.initial_y_ci95, self.initial_std_x, self.initial_std_y, self.initial_disp_std,  self.initial_disp_rmse, self.initial_disp_ci95 = self.get_summary_pt_stats_from_pd(self.points_pd)
        init_plot_name_x = os.path.join(self.analysis_folder, 'Displacement_BEFORE_X.png')
        init_plot_name_y = os.path.join(self.analysis_folder, 'Displacement_BEFORE_Y.png')
        print('Making preliminary histos...')
        self.make_fuzziness_histo(self.points_pd['Offset_X'], init_plot_name_x)
        self.make_fuzziness_histo(self.points_pd['Offset_Y'], init_plot_name_y)

        # create new PD with all the points adjusted
        # I've put this in its own function to keep it clean
        self.points_pd_ADJ = pd.DataFrame()
        self.post_summary_info = pd.DataFrame()
        self.make_adj_pts_pd()
        print('Calculating new statistics...')
        self.mean_x_centroid_offset_new = self.post_summary_info['New Offset X'].mean()
        self.mean_y_centroid_offset_new = self.post_summary_info['New Offset Y'].mean()

        self.new_x_rmse, self.new_x_ci95, self.new_y_rmse, self.new_y_ci95, self.new_std_x, self.new_std_y, self.new_disp_std, self.new_disp_rmse, self.new_disp_ci95 = self.get_summary_pt_stats_from_pd(self.points_pd_ADJ)
        new_plot_name_x = os.path.join(self.analysis_folder, 'Displacement_AFTER_X.png')
        new_plot_name_y = os.path.join(self.analysis_folder, 'Displacement_AFTER_Y.png')
        print('Making (hopefully) final histos...')
        self.make_fuzziness_histo(self.points_pd_ADJ['Offset_X'], new_plot_name_x)
        self.make_fuzziness_histo(self.points_pd_ADJ['Offset_Y'], new_plot_name_y)

        # Create various summary tables for report, make pages
        # total first
        print('Creating summary tables...')
        self.about_area_pd, self.adj_summary_pd, = self.make_summary_cover_tables()
        print('Making front page...')
        self.make_cover_page()

        # now per building
        for b in self.useful_buildings:
            # get relevant pts pd
            relevant_new_pts_pd = self.points_pd_ADJ[self.points_pd_ADJ['Building Name'] == b]
            relevant_old_pts_pd = self.points_pd[self.points_pd['Building Name'] == b]
            summary_info_table = self.get_per_building_info(relevant_new_pts_pd, relevant_old_pts_pd)
            # make doco
            self.make_sub_page(summary_info_table, b, relevant_old_pts_pd, relevant_new_pts_pd)
            print('Saved page for building %s...' % b)
        print('Complete...')

    def load_poly_file(self, poly_file):
        points_basename = os.path.basename(poly_file).replace('POLY_', 'POINTS_', 1).replace('_shp', '_pts')
        print('Running %s...' % points_basename)
        points_path = os.path.join(self.analysis_folder, points_basename)
        if not os.path.isfile(points_path):
            print(points_path)
            print(poly_file)
            print('error: points not found for ', os.path.basename(poly_file))
        else:
            self.poly_path_to_points_dic[poly_file] = points_path
            poly_poly = shp.Reader(poly_file)
            # 2. Load poly file
            #       2.1 check if it is an area or a line
            poly_check = True
            building_poly_coords = None
            building_poly = None
            try:
                building_poly_coords = poly_poly.shapeRecords()[0].shape.__geo_interface__['coordinates']
                building_string = LineString(building_poly_coords[0])
            except ValueError:
                print('linelist')
                line_list = []
                for a in poly_poly.shapeRecords():
                    building_poly_coords = a.shape.__geo_interface__['coordinates']
                    line_list.append(building_poly_coords)
                building_string = MultiLineString(line_list)
                poly_check = False
            building_buff = building_string.buffer(5)
            building_interior = building_buff.interiors
            # what was this meant to achieve??????
            try:
                building_poly = unary_union(building_string, building_interior)
            except Exception as e:
                # print('e3', e)
                print('unary exception')
                pass
            if poly_check:
                building_poly = Polygon(building_poly_coords[0])
            else:
                building_poly = building_string
            self.poly_shp_dic[poly_file] = (building_poly, building_string)

    def load_lidar_points(self, lidar_pts_file):
        points_list = []
        points_poly = shp.Reader(lidar_pts_file)
        for a in points_poly.shapeRecords():
            p = a.shape.__geo_interface__['coordinates']
            points_list.append(p)
        return points_list

    def find_pairs(self, building_name, mode):
        # return per-pt pd
        # Point, Building, Poly-Partner
        poly_file = self.name_to_poly[building_name]
        if 'prelim' in mode:
            lidar_pts_list = self.poly_to_points_list[poly_file]
        else:
            lidar_pts_list = self.poly_to_points_list_POST[poly_file]
        poly_pts = self.poly_shp_dic[poly_file][0]
        data_list = []
        headers = ['pair_id', 'Building Name', 'Point Number', 'Point', 'Point_X', 'Point_Y',
                   'Nearest Poly Pt', 'Nearest X', 'Nearest Y', 'Offset_X', 'Offset_Y', 'Displacement']
        for i, p in enumerate(lidar_pts_list):
            point_number = i
            p_as_Point = Point(p)
            nearest_poly = nearest_points(poly_pts, p_as_Point)[0]
            nearest_poly_x = nearest_poly.x
            nearest_poly_y = nearest_poly.y
            disp = nearest_poly.distance(p_as_Point)
            pair_id = '%s_%s' % (building_name, point_number)
            offset_x = nearest_poly_x - p[0]
            offset_y = nearest_poly_y - p[1]
            output_list = [pair_id, building_name, i, p, p[0], p[1], (nearest_poly_x, nearest_poly_y),
                           nearest_poly_x, nearest_poly_y, offset_x, offset_y, disp]
            data_list.append(output_list)
        output_pd = pd.DataFrame(columns=headers, data=data_list)
        return output_pd

    def get_ci_and_rmse(self, means_in):
        n = len(means_in)
        confidence = 0.95
        desired = [0] * n
        temp_rmse = round((np.sqrt(((np.array(desired) - np.array(means_in)) ** 2).mean())), 3)
        # m = mean(means_in)
        std_err = sem(means_in)
        h = std_err * t.ppf((1 + confidence) / 2, n - 1)
        # calced_ci = round((m + h), 3)
        temp_ci95 = 1.96 * temp_rmse
        return temp_ci95, temp_rmse

    def make_fuzziness_histo(self, distance_list_series, plot_name):
        sns.set(style="darkgrid")
        sns.set_color_codes()
        sns.distplot(distance_list_series.dropna(), norm_hist=True, color="r")

        plt.savefig(plot_name)
        sns.reset_defaults()
        sns.reset_orig()
        plt.clf()

    def make_adj_pts_pd(self):
        # return adj pd
        # step one: isolate the actual points
        working_pd = self.points_pd[['Building Name', 'Point_X', 'Point_Y']]
        # step two: shift the points

        def adj_x(row):
            x = row['Point_X']
            new_x = x + self.mean_x_centroid_offset
            return new_x

        def adj_y(row):
            y = row['Point_Y']
            new_y = y + self.mean_y_centroid_offset
            return new_y

        def adj_pt(row):
            x_adj = row['Point_X_Adj']
            y_adj = row['Point_Y_Adj']
            point_out = (x_adj, y_adj)
            return point_out

        working_pd['Point_X_Adj'] = working_pd.apply(adj_x, axis=1)
        working_pd['Point_Y_Adj'] = working_pd.apply(adj_y, axis=1)
        working_pd['Adjusted Point'] = working_pd.apply(adj_pt, axis=1)

        # Update poly to pts dict for analysis
        headers = ['Building Name', 'Pairs', 'New LiDAR Centroid',
                   'New Building Centroid', 'New Offset X', 'New Offset Y']
        per_building_summary_list_new = []
        all_useful_buildings = list(set(working_pd['Building Name'].to_list()))
        self.useful_buildings = all_useful_buildings
        for b in all_useful_buildings:
            poly_file = self.name_to_poly[b]
            relevant_pd = working_pd[working_pd['Building Name'] == b]
            pts_list = relevant_pd['Adjusted Point'].to_list()
            self.poly_to_points_list_POST[poly_file] = pts_list
            # get new pd
            new_pd = self.find_pairs(b, 'post')
            if self.points_pd_ADJ.empty:
                self.points_pd_ADJ = new_pd
            else:
                test = self.points_pd_ADJ
                self.points_pd_ADJ = pd.concat([test, new_pd], ignore_index=True, sort=False)

            # check new centroids
            lidar_points_as_MP = MultiPoint(pts_list)
            pair_points_as_MP = MultiPoint(new_pd['Nearest Poly Pt'].to_list())
            lidar_points_centroid = lidar_points_as_MP.centroid
            lidar_points_centroid_x = lidar_points_centroid.x
            lidar_points_centroid_y = lidar_points_centroid.y
            poly_points_centroid = pair_points_as_MP.centroid
            poly_points_centroid_x = poly_points_centroid.x
            poly_points_centroid_y = poly_points_centroid.y
            centroid_offset_x = poly_points_centroid_x - lidar_points_centroid_x
            centroid_offset_y = poly_points_centroid_y - lidar_points_centroid_y
            summary_info = [b, len(pts_list), (lidar_points_centroid_x, lidar_points_centroid_y),
                            (poly_points_centroid_x, poly_points_centroid_y), centroid_offset_x, centroid_offset_y]
            per_building_summary_list_new.append(summary_info)

        self.post_summary_info = pd.DataFrame(columns=headers, data=per_building_summary_list_new)

    def get_summary_pt_stats_from_pd(self, pts_pd_in):
        x_rmse, x_ci95 = self.get_ci_and_rmse(pts_pd_in['Offset_X'].to_list())
        y_rmse, y_ci95 = self.get_ci_and_rmse(pts_pd_in['Offset_Y'].to_list())
        std_x = pts_pd_in['Offset_X'].std()
        std_y = pts_pd_in['Offset_Y'].std()
        displacement_std = pts_pd_in['Displacement'].std()
        displacement_rmse, displacement_ci95 = self.get_ci_and_rmse(pts_pd_in['Displacement'].to_list())
        return x_rmse, x_ci95, y_rmse, y_ci95, std_x, std_y, displacement_std, displacement_rmse, displacement_ci95

    def make_summary_cover_tables(self):
        # about flight table
        # x, y offset, std table
        s_data = [['Site name', '%s %s' % (self.br, self.area_name)], ['Altitude flown (m)', self.flying_height],
                  ['Footprint (m)', self.ft], ['Building count', len(self.useful_buildings)],
                  ['LiDAR point count', len(self.points_pd_ADJ)]]
        s_pd = pd.DataFrame(data=s_data)

        x_series = pd.Series({'Offset': self.mean_x_centroid_offset, 'Standard Deviation after shift': self.new_std_x}, name='X')
        y_series = pd.Series({'Offset': self.mean_y_centroid_offset, 'Standard Deviation after shift': self.new_std_y}, name='Y')
        o_table = pd.concat([x_series, y_series], axis=1, ignore_index=False)

        return s_pd, o_table

    def build_doc(self, table_pd, template):
        doc = docxtpl.DocxTemplate(template)
        table_pd = table_pd.round(3)
        styles = doc.styles
        doco = doc.new_subdoc()
        doco_out = None
        pd.set_option('display.precision', 2)
        try:
            doco.add_paragraph("")
            tin = doco.add_table(table_pd.shape[0] + 1, table_pd.shape[1])
            tin.allow_autofit = True
            for j in range(table_pd.shape[-1]):
                tin.cell(0, j).text = str(table_pd.columns[j])
            for i in range(table_pd.shape[0]):
                for j in range(table_pd.shape[-1]):
                    tin.cell(i + 1, j).text = str(table_pd.values[i, j])
            for row in tin.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(9)
            doco_out = doco
        except KeyError:
            print('Could not gen table')
        return doco_out

    def determine_assessment(self):
        disp_before = self.points_pd['Displacement'].to_list()
        disp_after = self.points_pd_ADJ['Displacement'].to_list()

        t_test_statistic, t_test_p = ttest_rel(disp_before, disp_after, axis=0)

        if self.ft > self.initial_disp_std:
            assess = "As the footprint of %sm was more than the initial standard deviation of LiDAR point displacement" \
                     " of %sm, no horizontal adjustment has been applied to the data." % (self.ft,
                                                                                         round(self.initial_disp_std, 3))
        elif t_test_p < 0.05:
            assess = "As the change in total displacement RMSE was minimal (%s vs %s after, p-value of %s), no " \
                     "horizontal adjustment has been applied to the data." % (round(self.initial_disp_rmse, 3),
                                                                              round(self.new_disp_rmse, 3), t_test_p)
        elif self.ft < self.initial_disp_std:
            assess = "As the footprint of %sm is less than the initial standard deviation of %sm and there was a noted " \
                     "improvement in the displacement RMSE after testing the offset (%s vs %s after), a horizontal " \
                     "offset of %sm east, %sm north should be considered." % (self.ft, self.initial_disp_std,
                                                                              self.initial_disp_rmse,
                                                                              self.new_disp_rmse,
                                                                              self.mean_x_centroid_offset,
                                                                              self.mean_y_centroid_offset)
        else:
            assess = "Haha wow ok"
        return assess

    def make_cover_page(self):
        doc = docxtpl.DocxTemplate(front_page)
        date_now = datetime.today().strftime('%c')
        new_plot_name_x = os.path.join(self.analysis_folder, 'Displacement_AFTER_X.png')
        old_plot_name_x = os.path.join(self.analysis_folder, 'Displacement_BEFORE_X.png')
        new_plot_name_y = os.path.join(self.analysis_folder, 'Displacement_AFTER_Y.png')
        old_plot_name_y = os.path.join(self.analysis_folder, 'Displacement_BEFORE_Y.png')
        cover_path = os.path.join(self.reprots_folder, '%s_%s_Front_Page.docx' % (self.br, self.area_name))

        building_1_inline = docxtpl.InlineImage(doc, self.image_list[0], width=Mm(70))
        old_x_inline = docxtpl.InlineImage(doc, old_plot_name_x, width=Mm(70))
        old_y_inline = docxtpl.InlineImage(doc, old_plot_name_y, width=Mm(70))
        new_y_inline = docxtpl.InlineImage(doc, new_plot_name_y, width=Mm(70))
        new_x_inline = docxtpl.InlineImage(doc, new_plot_name_x, width=Mm(70))
        self.about_area_pd.infer_objects().style.set_precision(2)
        self.adj_summary_pd.infer_objects().style.set_precision(2)
        self.about_area_pd.columns = ['Variable', 'Value']

        ov_table = self.build_doc(self.about_area_pd, front_page)
        of_table = self.build_doc(self.adj_summary_pd, front_page)

        assessment = self.determine_assessment()

        context = {'BR': self.br,
                   'area_name': self.area_name,
                   'building_1': building_1_inline,
                   'overview_table': ov_table,
                   'offset_table': of_table,
                   'x_before': old_x_inline,
                   'y_before': old_y_inline,
                   'x_after': new_x_inline,
                   'y_after': new_y_inline,
                   'RMSE_X_before': self.initial_x_rmse,
                   'RMSE_X_after': self.new_x_rmse,
                   'RMSE_Y_before': self.initial_y_rmse,
                   'RMSE_Y_after': self.new_y_rmse,
                   'CI95_X_before': self.initial_x_ci95,
                   'CI95_X_after': self.new_x_ci95,
                   'CI95_Y_before': self.initial_y_ci95,
                   'CI95_Y_after': self.new_y_ci95,
                   'assessment': assessment,
                   'gen': date_now}
        doc.render(context)
        doc.save(cover_path)

    def get_per_building_info(self, new_pd, old_pd):
        relevant_new_pts_pd = new_pd
        # make relevant new mean, new ci95, new rmse, new std, histo x histo y pre post
        new_x_rmse, new_x_ci95, new_y_rmse, new_y_ci95, new_std_x, new_std_y, new_displacement_std, new_displacement_rmse, new_displacement_ci95 = self.get_summary_pt_stats_from_pd(
            relevant_new_pts_pd)
        new_x_mean = relevant_new_pts_pd['Offset_X'].mean()
        new_y_mean = relevant_new_pts_pd['Offset_Y'].mean()

        # get relevant old pts pd
        relevant_old_pts_pd = old_pd
        # ditto
        old_x_rmse, old_x_ci95, old_y_rmse, old_y_ci95, old_std_x, old_std_y, old_displacement_std, old_displacement_rmse, old_displacement_ci95 = self.get_summary_pt_stats_from_pd(
            relevant_old_pts_pd)
        old_x_mean = relevant_old_pts_pd['Offset_X'].mean()
        old_y_mean = relevant_old_pts_pd['Offset_Y'].mean()

        new_pd_out_headers = ['Value', 'Before adjustment', 'After adjustment']
        new_pd_data = [['X offset (m)', old_x_mean, new_x_mean], ['Y offset (m)', old_y_mean, new_y_mean],
                       ['X standard deviation (m)', old_std_x, new_std_x], ['Y standard deviation (m)', old_std_x, new_std_y],
                       ['X RMSE (m)', old_x_rmse, new_x_rmse], ['Y RMSE (m)', old_y_rmse, new_y_rmse],
                       ['X CI95', old_x_ci95, new_x_ci95], ['Y CI95', old_y_ci95, new_y_ci95]]
        new_pd_out = pd.DataFrame(columns=new_pd_out_headers, data=new_pd_data)
        new_pd_out.set_index('Value', inplace=True)

        return new_pd_out

    def make_sub_page(self, pd_in, building_name, old_data, new_data):
        page_number = old_data['Building ID'].head(1).values[0] + 2
        doc = docxtpl.DocxTemplate(sub_page)
        date_now = datetime.today().strftime('%c')
        page_path = os.path.join(self.reprots_folder, '%s_%s_%s_Page.docx' % (self.br, self.area_name, building_name))
        # make histos
        x_before_histo_path = os.path.join(self.analysis_folder, 'Displacement_%s_X_BEFORE.png' % building_name)
        y_before_histo_path = os.path.join(self.analysis_folder, 'Displacement_%s_Y_BEFORE.png' % building_name)
        y_after_histo_path = os.path.join(self.analysis_folder, 'Displacement_%s_Y_AFTER.png' % building_name)
        x_after_histo_path = os.path.join(self.analysis_folder, 'Displacement_%s_X_AFTER.png' % building_name)

        x_old_series = old_data['Offset_X']
        y_old_series = old_data['Offset_Y']
        y_new_series = new_data['Offset_Y']
        x_new_series = new_data['Offset_X']

        self.make_fuzziness_histo(x_old_series, x_before_histo_path)
        self.make_fuzziness_histo(y_old_series, y_before_histo_path)
        self.make_fuzziness_histo(y_new_series, y_after_histo_path)
        self.make_fuzziness_histo(x_new_series, x_after_histo_path)

        pd_in_reset_index = pd_in.reset_index().infer_objects()
        pd_in_reset_index.style.set_precision(precision=2)
        table_in = self.build_doc(pd_in_reset_index, sub_page)

        # find ss image
        longer_name = building_name + '.'
        candidates = [a for a in self.image_list if building_name in a]
        if len(candidates) == 1:
            building_image = candidates[0]
        elif len(candidates) > 1:
            candidates = [a for a in self.image_list if longer_name in a]
            building_image = candidates[0]
        else:
            building_image = None
        if building_image is not None:
            inline_image = docxtpl.InlineImage(doc, building_image, width=Mm(100))
        else:
            inline_image = "No image provided"

        x_before_histo_inline = docxtpl.InlineImage(doc, x_before_histo_path, width=Mm(70))
        y_before_histo_inline = docxtpl.InlineImage(doc, y_before_histo_path, width=Mm(70))
        y_after_histo_inline = docxtpl.InlineImage(doc, y_after_histo_path, width=Mm(70))
        x_after_histo_inline = docxtpl.InlineImage(doc, y_before_histo_path, width=Mm(70))

        context = {'BR': self.br,
                   'area_name': self.area_name,
                   'bnum': building_name,
                   'building_image': inline_image,
                   'building_adj_summary': table_in,
                   'x_histo_before': x_before_histo_inline,
                   'x_histo_after': x_after_histo_inline,
                   'y_histo_before': y_before_histo_inline,
                   'y_histo_after': y_after_histo_inline,
                   'RMSE_X_before': round(pd_in.loc['X RMSE (m)', 'Before adjustment'], 2),
                   'RMSE_X_after': round(pd_in.loc['X RMSE (m)', 'After adjustment'], 2),
                   'RMSE_Y_after': round(pd_in.loc['Y RMSE (m)', 'After adjustment'], 2),
                   'RMSE_Y_before': round(pd_in.loc['Y RMSE (m)', 'Before adjustment'], 2),
                   'CI95_Y_before': round(pd_in.loc['Y CI95', 'Before adjustment'], 2),
                   'CI95_X_before': round(pd_in.loc['X CI95', 'Before adjustment'], 2),
                   'CI95_X_after': round(pd_in.loc['X CI95', 'After adjustment'], 2),
                   'CI95_Y_after': round(pd_in.loc['Y CI95', 'After adjustment'], 2),
                   'gen': date_now,
                   'bid': page_number}
        doc.render(context)
        doc.save(page_path)


@Gooey(program_name="Horizontal Adjustments", use_legacy_titles=True, required_cols=1, default_size=(750, 500))
def goo():
    parser = GooeyParser(description='Validator')
    parser.add_argument('-af', metavar='Analysis folder', widget='DirChooser', help='Folder containing your shp and jpg for this process')
    parser.add_argument('-fh', metavar='Flying height (m)', help="Generic planned height is fine")
    parser.add_argument('-st', metavar='Sensor type', choices=['VQ780i', 'VQ780ii', 'H68'], default='VQ780ii')
    parser.add_argument('-br', metavar='Job BR code')
    parser.add_argument('-an', metavar='Area name, without BR code')
    return parser.parse_args()


tool(goo())
