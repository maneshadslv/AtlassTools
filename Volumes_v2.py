
from gooey import GooeyParser, Gooey
import pandas as pd
import json
import os
import subprocess
import matplotlib.path as mpltPath
import shapefile_old as sh
from multiprocessing import Pool, freeze_support, Manager
import time
import datetime
import shapefile_old as shp
from shapely.geometry import Point
from shapely.geometry.polygon import Polygon
from shapely.geometry import shape
from shapely.geometry.multipolygon import MultiPolygon
from shapely.geometry import box as shapely_box
from shapely.ops import unary_union

from decimal import Decimal
import docx
import docxtpl
import math
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import shutil

template_docx = r"Z:\PythonScripts\templates\volumes_report_template.docx"



zones = [str(a) for a in range(47, 58, 1)]
# print(zones)
map_sys = 'MGA_GDA94', 'MGA_GDA2020', 'AMG_AGD84', 'AMG_AGD66', 'ELL'


class prepare:
    def __init__(self, noot):
        print('###############################')
        print('Step 1 - Preparation - selected')
        print('###############################')
        print('')
        # sort variables
        self.noot = noot
        self.las_json = noot.las_json
        self.las_folder = noot.las_folder
        self.gm_exe = noot.gm_exe
        self.pre_shp_folder = noot.pre_shp_folder
        self.quick_name = noot.quick_name
        self.tile_size = int(noot.tile_size)
        self.shp_buff = noot.buffer_size
        self.map_system = noot.proj
        self.zone = noot.zone
        self.cores = int(noot.cores)
        self.open_it = noot.open_gmws
        self.problem_flag = []

        try:
            self.shp_list = [os.path.join(self.pre_shp_folder, a) for a in os.listdir(self.pre_shp_folder)
                             if a.endswith('.shp')]
        except TypeError:
            self.shp_list = []
        self.buff_shp_list = []

        # organise neighbours
        self.tile_pd = pd.DataFrame()
        self.sort_neighbours()

        # make buffered shp
        self.poly_pd = pd.DataFrame()
        self.define_areas()

        # make grids
        self.all_tiles = set()
        self.grids_folder = os.path.join(self.las_folder, "DEM_grids")
        self.grid_working = os.path.join(self.grids_folder, "working")
        if not os.path.isdir(self.grids_folder):
            os.mkdir(self.grids_folder)
        if not os.path.isdir(self.grid_working):
            os.mkdir(self.grid_working)
        self.grid_organiser()

        # make gmw with all contours, grids, las (as cat) loaded in
        self.build_world()
        print("All done!")
        if len(self.problem_flag) > 0:
            print('I did have issues with the following files though:')
            for a in self.problem_flag:
                print(a)
        else:
            print('No problem files were flagged at ctr/asc stage.')

    def sort_neighbours(self):
        print("Interpreting neighbours...")
        with open(self.las_json, 'r') as js:
            tl = json.load(js)

        for item in tl['features']:  # this may come in handy???
            cell_series = pd.Series(item['properties'], name=item['properties']['name'])
            temp = pd.concat([self.tile_pd, cell_series], axis=1, sort=True)
            self.tile_pd = temp
        self.tile_pd = self.tile_pd.T

        neighbours = []
        for item in tl['features']:
            p_dic = item['properties']
            sw_corner = (int(p_dic['xmin']), int(p_dic['ymin']))
            poss_x = []
            poss_y = []
            for variance in [(0-self.tile_size), 0, self.tile_size]:
                poss_x.append(int(sw_corner[0] + variance))
                poss_y.append(int(sw_corner[1] + variance))
            poss_new_sw = set()
            for xval in poss_x:
                for yval in poss_y:
                    if (xval, yval) != sw_corner:
                        poss_new_sw.add("%s_%s" % (xval, yval))
            neighbours.append(poss_new_sw)
        self.tile_pd['neighbours'] = neighbours

    def define_areas(self):
        print('Setting up stockpile zones...')
        # make buffered shapefiles
        for file in self.shp_list:
            self.buffer_shp(file)
        # work out which tiles are relevant
        x_points_list = self.tile_pd['xmin'].tolist()
        y_points_list = self.tile_pd['ymin'].tolist()
        name_list = self.tile_pd['name'].tolist()
        points_list = list(zip(x_points_list, y_points_list))
        if len(self.buff_shp_list) > 0:
            for file in self.buff_shp_list:
                # import it to something matplot will read
                listx = []
                listy = []
                fr = sh.Reader(file)
                for sr in fr.shapeRecords():
                    for x, y in sr.shape.points:
                        listx.append(x)
                        listy.append(y)
                poly = list(zip(listx, listy))
                path = mpltPath.Path(poly)
                inside_bools = path.contains_points(points_list)
                inside = []
                for i, tile in enumerate(name_list):
                    if inside_bools[i]:
                        inside.append(tile)
                file_info = {'file_name': file, 'coords': poly, 'tiles_inside': inside}
                poly_series = pd.Series(file_info, name=file)
                temp = pd.concat([self.poly_pd, poly_series], axis=1, sort=True)
                self.poly_pd = temp
        self.poly_pd = self.poly_pd.T

    def projection_definer(self):
        if 'AGD66' in self.map_system:
            proj_name = "AMG_ZONE%s_AUSTRALIAN_GEODETIC_1966" % self.zone
            proj_text = 'DEFINE_PROJ PROJ_NAME="AMG_ZONE%s_AUSTRALIAN_GEODETIC_1966"\n' \
                        'Projection     AMG (Australian Map Grid)\n' \
                        'Datum          D_AUSTRALIAN_1966\n' \
                        'Zunits         NO\n' \
                        'Units          METERS\n' \
                        'Zone           %s\n' \
                        'Xshift         0.000000\n' \
                        'Yshift         0.000000\n' \
                        'Parameters\n' \
                        'END_DEFINE_PROJ' % (self.zone, self.zone)
        elif 'AGD84' in self.map_system:
            proj_name = "AMG_ZONE%s_AUSTRALIAN_GEODETIC_1984" % self.zone
            proj_text = 'DEFINE_PROJ PROJ_NAME="AMG_ZONE%s_AUSTRALIAN_GEODETIC_1984"\n' \
                        'Projection     AMG (Australian Map Grid)\n' \
                        'Datum          D_AUSTRALIAN_1984\n' \
                        'Zunits         NO\n' \
                        'Units          METERS\n' \
                        'Zone           %s\n' \
                        'Xshift         0.000000\n' \
                        'Yshift         0.000000\n' \
                        'Parameters\n' \
                        'END_DEFINE_PROJ' % (self.zone, self.zone)
        elif 'GDA2020' in self.map_system:
            proj_name = "MGA_ZONE%s_GDA_2020_AUSTRALIAN_GEODETIC_2020" % self.zone
            proj_text = 'DEFINE_PROJ PROJ_NAME="MGA_ZONE%s_GDA_2020_AUSTRALIAN_GEODETIC_2020"\n' \
                        'Projection     MGA (Map Grid of Australia)\n' \
                        'Datum          GDA2020\n' \
                        'Zunits         NO\n' \
                        'Units          METERS\n' \
                        'Zone           %s\n' \
                        'Xshift         0.000000\n' \
                        'Yshift         0.000000\n' \
                        'Parameters\n' \
                        'END_DEFINE_PROJ' % (self.zone, self.zone)
        elif 'GDA94' in self.map_system:
            proj_name = "MGA_ZONE%s_GDA_94_AUSTRALIAN_GEODETIC_1994" % self.zone
            proj_text = 'DEFINE_PROJ PROJ_NAME="MGA_ZONE%s_GDA_94_AUSTRALIAN_GEODETIC_1994"\n' \
                        'Projection     MGA (Map Grid of Australia)\n' \
                        'Datum          GDA94\n' \
                        'Zunits         NO\n' \
                        'Units          METERS\n' \
                        'Zone           %s\n' \
                        'Xshift         0.000000\n' \
                        'Yshift         0.000000\n' \
                        'Parameters\n' \
                        'END_DEFINE_PROJ' % (self.zone, self.zone)
        else:
            print("Sorry, I don't do that yet...")
            proj_name = None
            proj_text = None
        return proj_text, proj_name

    def buffer_shp(self, in_file):
        buff = int(self.shp_buff)
        if 'buffed' not in in_file:
            temp_gms_name = os.path.splitext(in_file)[0] + "_buffing.gms"
            outfile_name = os.path.splitext(in_file)[0] + "_buffed.shp"
            if os.path.isfile(temp_gms_name):
                os.remove(temp_gms_name)
            if os.path.isfile(outfile_name):
                os.remove(outfile_name)
                for ext in ['.shx', '.prj', '.dbf']:
                    f = os.path.splitext(in_file)[0] + "_buffed" + ext
                    os.remove(f)
            projection_string, proj_name = self.projection_definer()
            gms_str = "GLOBAL_MAPPER_SCRIPT VERSION=1.00\n" \
                      "%s\n" \
                      "IMPORT FILENAME=%s TYPE=AUTO PROJ=%s\n" \
                      "EDIT_VECTOR FILENAME=* SHAPE_TYPE=AREAS BUFFER_DIST=%s BUFFER_COMBINE_AREAS=YES " \
                      "NEW_LAYER_NAME='Buffered'\n" \
                      "EXPORT_VECTOR EXPORT_LAYER='Buffered' TYPE=SHAPEFILE FILENAME=%s SHAPE_TYPE=AREAS GEN_PRJ_FILE=YES" \
                      " OVERWRITE_EXISTING=YES\n" \
                      "UNLOAD_ALL" % (projection_string, in_file, proj_name, buff, outfile_name)
            with open(temp_gms_name, 'w') as wg:
                wg.write(gms_str)
            command = [self.gm_exe, temp_gms_name]
            print('Running command to buffer shapefile...')
            subprocess.call(command)
            self.buff_shp_list.append(outfile_name)
        else:
            print('Disregarding exisiting file', os.path.basename(in_file))

    def grid_organiser(self):
        print('Building grids...')
        if len(self.buff_shp_list) < 0:
            for polygon in self.buff_shp_list:
                inside_tiles = self.poly_pd.loc[[polygon], ['tiles_inside']].values[0][0]
                for a in inside_tiles:
                    self.all_tiles.add(a)
        else:
            self.all_tiles = self.tile_pd['name'].to_list()
        self.sanity_check()

    def crop_files(self, tile, folder):
        t_b = tile
        xmin = str(self.tile_pd.loc[[t_b], ['xmin']].values[0][0])
        xmax = str(self.tile_pd.loc[[t_b], ['xmax']].values[0][0])
        ymin = str(self.tile_pd.loc[[t_b], ['ymin']].values[0][0])
        ymax = str(self.tile_pd.loc[[t_b], ['ymax']].values[0][0])

        pr_text, pr_name = self.projection_definer()

        input_dem = os.path.join(folder, t_b + "_m.asc")
        output_dem = os.path.join(folder, t_b + "_out.asc")
        input_ctr = os.path.join(folder, t_b + "_m.shp")
        output_ctr = os.path.join(folder, t_b + "_out.shp")

        crop_gms = os.path.join(folder, "%s_crop.gms" % t_b)

        crop_gms_str = 'GLOBAL_MAPPER_SCRIPT VERSION="1.00"\n' \
                       'UNLOAD_ALL\n' \
                       '%s\n' \
                       'IMPORT FILENAME=%s ELEV_UNITS=METERS PROJ=%s\n' \
                       'IMPORT FILENAME=%s ELEV_UNITS=METERS PROJ=%s\n' \
                       'EXPORT_VECTOR FILENAME=%s EXPORT_LAYER=%s GEN_PRJ_FILE=YES SHAPE_TYPE=LINES GLOBAL_BOUNDS="%s,%s,%s,%s" TYPE=SHAPEFILE\n' \
                       'EXPORT_ELEVATION FILENAME=%s EXPORT_LAYER=%s GEN_PRJ_FILE=YES GLOBAL_BOUNDS="%s,%s,%s,%s" TYPE=ARCASCIIGRID\n' \
                       'UNLOAD_ALL' % (pr_text, input_dem, pr_name, input_ctr, pr_name, output_ctr, input_ctr, xmin,
                                       ymin, xmax, ymax, output_dem, input_dem, xmin, ymin, xmax, ymax)

        with open(crop_gms, 'w') as cg:
            cg.write(crop_gms_str)
        cg_str = [self.gm_exe, crop_gms]
        subprocess.call(cg_str)

    def file_maker(self, tile):
        # make folder in working for tile
        if not tile.endswith('.laz'):
            tile = tile + ".laz"
        tile = os.path.join(self.las_folder, tile)
        t_b = os.path.splitext(os.path.basename(tile))[0]  # this makes sense later
        working_folder = os.path.join(self.grid_working, t_b)
        if not os.path.isdir(working_folder):
            os.mkdir(working_folder)
        # get neighbours
        neighbours = list(self.tile_pd.loc[[t_b], ['neighbours']].values[0][0])

        # merge w neighbours in working/tile
        files_to_merge = [os.path.join(self.las_folder, a + '.laz') for a in neighbours]
        files_to_merge.append(tile)
        out_laz_name = t_b + "_m.laz"
        out_laz_path = os.path.join(working_folder, out_laz_name)
        las_merge_command = ['las2las', '-i']
        for a in files_to_merge:
            las_merge_command.append(a)
        las_merge_command.extend(['-merged', '-olaz', '-o', out_laz_path])
        dem_command = ['blast2dem', '-i', out_laz_path, '-step', '0.5', '-keep_class', '2', '-oasc']
        ctr_command = ['las2iso', '-i', out_laz_path, '-keep_class', '2', '-odbf', '-oshp', '-iso_every', '0.1']

        # run
        subprocess.call(las_merge_command)
        # subprocess.call(las_info_bb_command)
        subprocess.call(dem_command)
        subprocess.call(ctr_command)

        # clip
        asc_name = out_laz_name.replace('.laz', '.asc')
        ctr_name = out_laz_name.replace('.laz', '.shp')
        self.crop_files(t_b, working_folder)

        asc_path = os.path.join(working_folder, asc_name)
        ctr_path = os.path.join(working_folder, ctr_name)
        new_asc_name = asc_path.replace('.asc', '_old_asc')
        new_ctr_name = ctr_path.replace('.shp', '_old_ctr')

        try:
            os.replace(asc_path, new_asc_name)  # clumsy but whatever
            os.replace(ctr_path, new_ctr_name)
        except FileNotFoundError:
            self.problem_flag.append(asc_name)

        for f in os.listdir(working_folder):
            if not f.endswith('.laz'):
                if not f.endswith('_old_asc'):
                    if not f.endswith('_old_ctr'):
                        f_p = os.path.join(working_folder, f)
                        f_new_p = os.path.join(self.grids_folder, f)
                        os.replace(f_p, f_new_p)

    def sanity_check(self):
        freeze_support()
        print('Running file prep...')
        pool = Pool(processes=self.cores)
        pool.map(self.file_maker, list(self.all_tiles))
        pool.close()
        pool.join()

    def build_world(self):
        print('Building world...')
        world_name = "Volumes_%s_world.gmw" % self.quick_name
        world_path = os.path.join(self.grids_folder, world_name)
        dem_cat = os.path.join(self.grids_folder, "dem_files.gmc")
        dem_filter = "%s\\*.asc" % self.grids_folder
        ctr_cat = os.path.join(self.grids_folder, 'ctr_files.gmc')
        ctr_filter = "%s\\*.shp" % self.grids_folder
        pr_text, pr_name = self.projection_definer()
        world_make_gms = os.path.join(self.grids_folder, "worldmaker.gms")
        world_maker_str = 'GLOBAL_MAPPER_SCRIPT VERSION="1.00"\n' \
                          'UNLOAD_ALL\n' \
                          '%s\n' \
                          'EDIT_MAP_CATALOG FILENAME=%s CREATE_IF_EMPTY=YES ADD_FILE=%s ZOOM_DISPLAY="PERCENT,0.90,0" PROJ=%s\n' \
                          'EDIT_MAP_CATALOG FILENAME=%s CREATE_IF_EMPTY=YES ADD_FILE=%s ZOOM_DISPLAY="PERCENT,0.90,0" PROJ=%s\n' \
                          'IMPORT FILENAME=%s ELEV_UNITS=METERS\n' \
                          'IMPORT FILENAME=%s ELEV_UNITS=METERS\n' \
                          'SAVE_WORKSPACE FILENAME=%s\n' \
                          'UNLOAD_ALL' % (pr_text, dem_cat, dem_filter, pr_name, ctr_cat, ctr_filter, pr_name, dem_cat, ctr_cat, world_path)
        with open(world_make_gms, 'w') as wmg:
            wmg.write(world_maker_str)
        gmw_make_str = [self.gm_exe, world_make_gms]
        subprocess.call(gmw_make_str)
        if self.open_it:
            gmw_comm = [self.gm_exe, world_path]
            subprocess.Popen(gmw_comm, shell=False, stdin=None, stdout=None, stderr=None, close_fds=True)


class calculate:
    def __init__(self, noot):
        print('###############################')
        print('Step 2 - Calculation - selected')
        print('###############################')
        print('')
        self.time_since_epoch = 315964782
        self.info_dic = {}
        self.times_dic = {}
        self.pd_li_dic = {}
        self.toe_information = {}

        self.cores = noot.cores
        self.world = noot.gmw_file
        self.project_name = noot.name.replace(' ', '_')
        self.dem_folder = noot.dem_folder
        self.map_system = noot.proj
        self.zone = noot.zone
        self.client_name = noot.client_name
        self.client_company = noot.client_company
        self.user_name = noot.user_name
        self.product_type = noot.product_type

        self.base_folder = noot.base_folder
        self.base_list = []
        self.base_cat = None
        if self.base_folder is not None:
            if os.path.isdir(self.base_folder):
                self.base_list = [os.path.join(self.base_folder, a) for a in os.listdir(self.base_folder)
                                  if a.endswith('.gmg')]
        
            self.base_cat = os.path.join(self.base_folder, "base_files.gmc")

        self.toes_per_area_files = []
        self.toe_area_database = pd.DataFrame()

        self.laz_folder = noot.laz_folder
        self.laz_cat = os.path.join(self.laz_folder, 'laz_files.gmc')
        self.toes_folder = os.path.join(self.laz_folder, 'exported_toes')
        self.toes_per_area_folder = os.path.join(self.toes_folder, 'Areas')
        self.csv_folder = os.path.join(self.toes_folder, 'csv_reports')
        if not os.path.isdir(self.toes_folder):
            os.mkdir(self.toes_folder)
        self.gm_exe = noot.gm_exe
        self.step_3 = noot.include_step_3
        self.error_log = os.path.join(self.csv_folder, "errors.txt")

        self.template_file = os.path.join(self.laz_folder, 'template.gmw')

        # pandas
        self.toes_times_pd = pd.DataFrame()

        # misc dics
        self.toe_to_base = {}
        self.base_not_available = []

        # if this is false, it drapes the toes
        self.toes_preset = noot.toe_technique

        # define some GMS
        self.toe_prepper_gms = os.path.join(self.laz_folder, "toe_prepper.gms")
        self.laz_gmc = os.path.join(self.laz_folder, "laz_files.gmc")

        self.csv_dic = {}
        self.final_csv = os.path.join(self.csv_folder, '%s_FINAL.csv' % self.project_name)

        # BETA
        self.output_xl_path = os.path.join(self.csv_folder, 'FINAL.xlsx')

        # Processes
        self.toe_prepper()

        self.toes_list = [os.path.join(self.toes_folder, a) for a in os.listdir(self.toes_folder)
                          if a.endswith('.shp')]
        self.build_template_world()
        self.world_running()
        self.csv_tool()

        if self.step_3:
            report_dic = {'project_name': self.project_name,
                          'csv folder': self.csv_folder, 'user_name': self.user_name, 'client_name': self.client_name,
                          'client_company': self.client_company, 'product_type': self.product_type,
                          'show_all_vals': False}
            report_make(report_dic)
        else:
            print('Kindly now review your CSV output.')

    def projection_definer(self):
        if 'AGD66' in self.map_system:
            proj_name = "AMG_ZONE%s_AUSTRALIAN_GEODETIC_1966" % self.zone
            proj_text = 'DEFINE_PROJ PROJ_NAME="AMG_ZONE%s_AUSTRALIAN_GEODETIC_1966"\n' \
                        'Projection     AMG (Australian Map Grid)\n' \
                        'Datum          D_AUSTRALIAN_1966\n' \
                        'Zunits         NO\n' \
                        'Units          METERS\n' \
                        'Zone           %s\n' \
                        'Xshift         0.000000\n' \
                        'Yshift         0.000000\n' \
                        'Parameters\n' \
                        'END_DEFINE_PROJ' % (self.zone, self.zone)
        elif 'AGD84' in self.map_system:
            proj_name = "AMG_ZONE%s_AUSTRALIAN_GEODETIC_1984" % self.zone
            proj_text = 'DEFINE_PROJ PROJ_NAME="AMG_ZONE%s_AUSTRALIAN_GEODETIC_1984"\n' \
                        'Projection     AMG (Australian Map Grid)\n' \
                        'Datum          D_AUSTRALIAN_1984\n' \
                        'Zunits         NO\n' \
                        'Units          METERS\n' \
                        'Zone           %s\n' \
                        'Xshift         0.000000\n' \
                        'Yshift         0.000000\n' \
                        'Parameters\n' \
                        'END_DEFINE_PROJ' % (self.zone, self.zone)
        elif 'GDA2020' in self.map_system:
            proj_name = "MGA_ZONE%s_GDA_2020_AUSTRALIAN_GEODETIC_2020" % self.zone
            proj_text = 'DEFINE_PROJ PROJ_NAME="MGA_ZONE%s_GDA_2020_AUSTRALIAN_GEODETIC_2020"\n' \
                        'Projection     MGA (Map Grid of Australia)\n' \
                        'Datum          GDA2020\n' \
                        'Zunits         NO\n' \
                        'Units          METERS\n' \
                        'Zone           %s\n' \
                        'Xshift         0.000000\n' \
                        'Yshift         0.000000\n' \
                        'Parameters\n' \
                        'END_DEFINE_PROJ' % (self.zone, self.zone)
        elif 'GDA94' in self.map_system:
            proj_name = "MGA_ZONE%s_GDA_94_AUSTRALIAN_GEODETIC_1994" % self.zone
            proj_text = 'DEFINE_PROJ PROJ_NAME="MGA_ZONE%s_GDA_94_AUSTRALIAN_GEODETIC_1994"\n' \
                        'Projection     MGA (Map Grid of Australia)\n' \
                        'Datum          GDA94\n' \
                        'Zunits         NO\n' \
                        'Units          METERS\n' \
                        'Zone           %s\n' \
                        'Xshift         0.000000\n' \
                        'Yshift         0.000000\n' \
                        'Parameters\n' \
                        'END_DEFINE_PROJ' % (self.zone, self.zone)
        else:
            print("Sorry, I don't do that yet...")
            proj_name = None
            proj_text = None
        return proj_text, proj_name

    def toe_prepper(self):
        print("Preparing toes for processing...")
        pr_text, pr_name = self.projection_definer()
        # todo: make this have a log as its error message is bloody annoying
        tp_string_1 = 'GLOBAL_MAPPER_SCRIPT VERSION=1.00 ENABLE_PROGRESS=YES LOG_TO_COMMAND_PROMPT="YES"\n' \
                      'IMPORT FILENAME=%s ELEV_UNITS=METERS\n' \
                      '%s\n' \
                      'LAYER_LOOP_START FILENAME="Toes*"\n\t' \
                      % (self.world, pr_text)
        if not self.toes_preset:
            tp_string_1 = tp_string_1 + 'EDIT_VECTOR FILENAME=%LAYER_DESC% APPLY_ELEVS=YES REPLACE_EXISTING=YES\n\t'
        tp_string_2 = tp_string_1 + 'EXPORT_VECTOR FILENAME=%s/%s_.shp EXPORT_LAYER=%s ' \
                                    'GEN_PRJ_FILE=YES SHAPE_TYPE=AREAS SPLIT_BY_ATTR=YES ' \
                                    'FILENAME_ATTR="<Feature Name>" TYPE=SHAPEFILE INC_ELEV_ATTR=YES ' \
                                    'INC_LAYER_ATTR=YES GEN_3D_FEATURES=YES INC_MAP_NAME_ATTR=YES ' \
                                    'INC_STYLE_ATTRS=YES\n\t' \
                                    % (self.toes_folder, '%LAYER_DESC%', '%LAYER_DESC%')
        if not os.path.isfile(self.laz_gmc):
            tp_string_3 = tp_string_2 + 'EDIT_MAP_CATALOG FILENAME=%s CREATE_IF_EMPTY=YES ADD_FILE=%s/*.laz PROJ=%s\n\t'\
                                        % (self.laz_gmc, self.laz_folder, pr_name)
        else:
            tp_string_3 = tp_string_2
        tp_string_4 = tp_string_3 + 'IMPORT FILENAME=%s\n\t' \
                                    'EXPORT_VECTOR EXPORT_LAYER=%s TYPE=LIDAR_LAS FILENAME=%s/cut_.laz ' \
                                    'POLYGON_CROP_FILE=%s POLYGON_CROP_USE_EACH=YES ' \
                                    'POLYGON_CROP_NAME_ATTR="<Feature Name>"\n\t' % (self.laz_gmc, self.laz_gmc,
                                                                                     self.toes_folder, '%LAYER_DESC%')
        # export for quick list of toes we need
        tp_string_5 = tp_string_4 + 'EXPORT_VECTOR EXPORT_LAYER=%s TYPE=SHAPEFILE FILENAME=%s/%s.shp ' \
                                    'SHAPE_TYPE=AREAS\n LAYER_LOOP_END' \
                                    % ('%LAYER_DESC%', self.toes_per_area_folder, '%LAYER_DESC%')

        with open(self.toe_prepper_gms, 'w') as toe:
            toe.write(tp_string_5)
        subprocess.call([self.gm_exe, self.toe_prepper_gms])
        self.toe_time_processor()

    def toe_time_processor(self):
        print("Calculating toe times...")
        toes_files = [os.path.join(self.toes_folder, a) for a in os.listdir(self.toes_folder) if a .endswith('.laz')]
        if len(toes_files) > 0:
            self.pool_boy_lasinfo(toes_files)
        else:
            print('No toes found')
            exit()
        clean_dict = {}
        for item in self.times_dic.keys():
            item_bname = os.path.splitext(os.path.basename(item))[0].replace('cut_', '')
            clean_dict[item_bname] = self.times_dic[item]
        self.toes_times_pd = pd.DataFrame.from_dict(clean_dict, orient='columns')

    def gps_to_utc(self, test):
        test = test + self.time_since_epoch
        test2 = time.strftime("%b %d %Y %H:%M:%S", time.gmtime(test))
        return test2

    def lasinfo_tool(self, file_in):
        file_base = os.path.splitext(os.path.basename(file_in))[0]
        text_out = os.path.splitext(file_in)[0] + '.txt'
        line_vals = {'file_name': file_in}
        string_in = ['lasinfo', '-i', file_in, '-nh', '-nr', '-nw']
        run_command = (subprocess.check_output(string_in, stderr=subprocess.STDOUT, shell=True)).decode('UTF-8')
        time_line = run_command[run_command.find('  gps_time '):run_command.rfind('  Color')]
        line_vals['latest_time'] = time_line.split()[1].strip()
        line_vals['human_time'] = self.gps_to_utc(float(time_line.split()[1].strip()))
        self.times_dic[file_in] = line_vals
        with open(text_out, 'w') as tow:
            tow.write(str(line_vals))
            tow.close()
        file_series = pd.Series(line_vals, name=file_in)
        self.pd_li_dic[file_base] = file_series

    def pool_boy_lasinfo(self, list_in):
        print('poolboy_found')
        if __name__ == '__main__':
            freeze_support()
            manager = Manager()
            self.times_dic = manager.dict()

            pool = Pool(processes=int(self.cores))
            pool.map(self.lasinfo_tool, list_in)
            pool.close()
            pool.join()
        print('poolboy done')

    def build_template_world(self):
        print('Building template world...')
        #   - base files location (self.base_folder)
        # first, copy original gmw. remove all layers not cats, then save as template
        # also set shader to daylight shader
        # NEED LAS CAT
        pr_text, pr_name = self.projection_definer()
        script_string = 'GLOBAL_MAPPER_SCRIPT VERSION=1.00 SHOW_WARNINGS=NO\n' \
                        'UNLOAD_ALL\n' \
                        'IMPORT FILENAME=%s\n' \
                        'LAYER_LOOP_START FILENAME="*"\n' \
                        'IF COMPARE_STR=%s!=dem_files.gmc\n' \
                        '\tUNLOAD_LAYER FILENAME=%s\n' \
                        'ELSE\n' \
                        '\tSET_LAYER_OPTIONS FILENAME="%s" SHADER_NAME="Daylight Shader"\n' \
                        'END_IF\n' \
                        'LAYER_LOOP_END\n' % (self.world, '%LAYER_FNAME%', '%LAYER_DESC%', '%LAYER_DESC%')
        if not os.path.isfile(self.laz_gmc):
            script_string = script_string + 'EDIT_MAP_CATALOG FILENAME=%s CREATE_IF_EMPTY=YES ADD_FILE=%s/*.laz ' \
                                            'ZOOM_DISPLAY="PERCENT,0.90,0" PROJ=%s\n' % (self.laz_gmc, self.laz_folder,
                                                                                         pr_name)
        
        if self.base_folder is not None:
            if not os.path.isfile(self.base_cat):
                script_string = script_string + 'EDIT_MAP_CATALOG FILENAME=%s CREATE_IF_EMPTY=YES ADD_FILE=%s/*.gmg ' \
                                            'ZOOM_DISPLAY="PERCENT,0.90,0" PROJ=%s\n' % (self.base_cat, self.base_folder,
                                                                                         pr_name)
        if self.base_folder is not None:
            script_string = script_string + 'IMPORT FILENAME=%s\n' \
                                            'IMPORT FILENAME=%s\n' \
                                            'SAVE_WORKSPACE FILENAME=%s\n' \
                                            'UNLOAD_ALL' % (self.laz_gmc, self.base_cat, self.template_file)
        else:
            script_string = script_string + 'IMPORT FILENAME=%s\n' \
                                            'SAVE_WORKSPACE FILENAME=%s\n' \
                                            'UNLOAD_ALL' % (self.laz_gmc, self.template_file)

        template_gms_name = os.path.join(self.laz_folder, 'template_build.gms')
        with open(template_gms_name, 'w') as tg:
            tg.write(script_string)
        subprocess.call([self.gm_exe, template_gms_name])

    def world_running(self):
        # Actually works best with toes info run first
        print('Working out base coverage...')
        # todo: change to use shapely
        for base in self.base_list:
            print('base list', self.base_list)
            temp_gms = os.path.join(self.base_folder, 'temp.gms')
            temp_output = os.path.splitext(base)[0] + "_bounding_box.shp"
            temp_output_laz = os.path.splitext(base)[0] + "_bounding_box.laz"
            temp_gms_str = "GLOBAL_MAPPER_SCRIPT VERSION=1.00 SHOW_WARNINGS=NO\n" \
                           "UNLOAD_ALL\n" \
                           "IMPORT FILENAME=%s TYPE='AUTO'\n" \
                           "EXPORT_ELEVATION FILENAME=%s TYPE=LIDAR_LAS" % (base, temp_output_laz)
            with open(temp_gms, 'w') as tg:
                tg.write(temp_gms_str)
            subprocess.call([self.gm_exe, temp_gms])
            lasboundary_str = ['lasboundary', '-i', str(temp_output_laz), '-oshp', '-o', str(temp_output),
                               '-concavity', '5', '-holes', '-disjoint']
            subprocess.call(lasboundary_str)

        base_bounds_list = [os.path.join(self.base_folder, b) for b in os.listdir(self.base_folder) if
                            b.endswith('.shp')]
        if self.base_folder is not None:
            base_bounds_list = [os.path.join(self.base_folder, b) for b in os.listdir(self.base_folder) if
                                b.endswith('.shp')]
        for toe in self.toes_list:
            toe_coverage = []
            toe_partial_coverage = []
            for base in base_bounds_list:
                # check if base covers toe
                # first, import toe shp as shapefile. If toe is covered, all its points are in the base area.
                # if >80% (arbitrary) pts covered, data is probably ok, but will need to be flagged.
                toe_read = shp.Reader(toe)
                toe_points_list = toe_read.shape(0).points
                base_read_pyshp = shp.Reader(base)
                base_areas = []
                for c in base_read_pyshp.shapes():
                    coords = c.__geo_interface__['coordinates'][0]
                    coords_as_list = list(coords)
                    base_areas.append(Polygon(coords_as_list))
                base_Poly_mode = unary_union(base_areas)

                points_in = []
                for point in toe_points_list:
                    p = Point(point)
                    if p.within(base_Poly_mode):
                        points_in.append(p)
                if len(points_in) > 0:
                    cov_ratio = len(toe_points_list) / len(points_in)
                else:
                    cov_ratio = 0
                if cov_ratio == 1.0:
                    toe_coverage.append(base)
                elif cov_ratio > 0.8:
                    toe_partial_coverage.append(base)
                else:
                    pass
            if len(toe_coverage) == 0 and len(toe_partial_coverage) == 0:
                self.base_not_available.append(toe)
            base_name = os.path.splitext(os.path.basename(toe))[0]
            self.toe_information[base_name] = {'base_coverage': toe_coverage, 'partial': toe_partial_coverage,
                                               'file_path': toe}
        for t in self.toe_information.keys():
            if len(self.toe_information[t]['base_coverage']) == 0:
                if len(self.toe_information[t]['partial']) == 0:
                    self.base_not_available.append(t)
        # print("BNA", self.base_not_available)
        # print("toe info", self.toe_information)
        # loads up per rough spile "area" BUT NOW GENERATES THESE FROM THE ACTUAL TOES
        self.toes_per_area_files = [os.path.join(self.toes_per_area_folder, a) for a in
                                    os.listdir(self.toes_per_area_folder) if a.endswith('.shp')]
        # quickly generate a database listing which toes belong in which area,
        # and buffered bounds for their load into Global Mapper
        print('Calculating volumes!')
        for area in self.toes_per_area_files:
            # we want to see the listings of poly names, first up

            area_read = shp.Reader(area)
            quick_area_list = []
            for spile_area in area_read.records():
                spile_area_series = pd.Series(spile_area.as_dict(), name=spile_area['NAME']).T
                temp = self.toe_area_database.append(spile_area_series, ignore_index=False)
                self.toe_area_database = temp
                quick_area_list.append(spile_area['NAME'])
            # that can safely compile without amendment; we can recover "area" name from ['BNDRY_INFO'] later
            # now we want bounding box, so we can buffer appropriately in screenshot
            xmin, ymin, xmax, ymax = area_read.bbox
            bbx_poly = shapely_box(xmin, ymin, xmax, ymax)
            buff_area = bbx_poly.buffer(200, resolution=16, cap_style=3, join_style=1, mitre_limit=5.0)
            buff_xmin, buff_ymin, buff_xmax, buff_ymax = buff_area.bounds

            area_base = os.path.splitext(os.path.basename(area))[0]
            print('Calculating for %s...' % area_base)
            layer_spile_list_paths = [os.path.join(self.toes_folder, "%s_%s.shp" % (area_base, a)) for a in quick_area_list]
            layer_spile_list = ["%s_%s" % (area_base, a) for a in quick_area_list]

            # now to make the world!
            # we need:
            #   - template file name (has laz, dem, bases preloaded)
            #   - world name (area_base)
            #   - view value (buff_xmin, buff_ymin, buff_xmax, buff_ymax)
            #   - csv output location (self.csv_folder)
            #   - list of relevant toes (layer_spile_list)
            #   - which ones have no base coverage (toe_coverage)

            # open template world. import covered toes as layer group "COVERED". calc s2s vols.
            # import uncovered toes. create a grid made from these toes. calc all vols.
            # export all to pdf to extents in buff_area
            # save world to area_base.gmw
            pr_text, pr_name = self.projection_definer()
            output_t2t = os.path.join(self.csv_folder, '%s_T2T.csv' % area_base)
            output_s2s = os.path.join(self.csv_folder, '%s_S2S.csv' % area_base)
            output_pdf_page = os.path.join(self.csv_folder, '%s_page.pdf' % area_base)
            output_workspace = os.path.join(self.csv_folder, '%s_World.gmw' % area_base)
            gms_name = os.path.join(self.csv_folder, '%s_Worldbuilder.gms' % area_base)

            self.csv_dic[area_base] = {'T2T': output_t2t, 'S2S': output_s2s, 'PDF': output_pdf_page}

            start_string = 'GLOBAL_MAPPER_SCRIPT VERSION=1.00 SHOW_WARNINGS=NO\n' \
                           'UNLOAD_ALL\n' \
                           'IMPORT FILENAME=%s\n' % self.template_file
            import_covered_string = ""
            not_yet_imported = []
            # print(self.toe_information.keys())
            base_yn = False
            nc_flag = False
            for i, toe in enumerate(layer_spile_list):
                toe_info = self.toe_information[toe]
                print('tibc', toe_info['partial'])
                toe_path = layer_spile_list_paths[i]
                if len(toe_info['base_coverage']) > 0:
                    this_toe = 'IMPORT FILENAME=%s TYPE=SHAPEFILE LAYER_GROUP="POLY_COVERED" PROJ=%s\n' \
                               % (toe_path, pr_name)
                    import_covered_string = import_covered_string + this_toe
                    base_yn = True
                elif len(toe_info['partial']) > 0:
                    this_toe = 'IMPORT FILENAME=%s TYPE=SHAPEFILE LAYER_GROUP="POLY_COVERED" PROJ=%s\n' \
                               % (toe_path, pr_name)
                    import_covered_string = import_covered_string + this_toe
                    base_yn = True
                else:
                    this_toe = 'IMPORT FILENAME=%s TYPE=SHAPEFILE LAYER_GROUP="POLY_NOT_COVERED" PROJ=%s\n' \
                               % (toe_path, pr_name)
                    import_covered_string = import_covered_string + this_toe
                    nc_flag = True

            # move polygons to single layer so calc volumes will take them
            copy_string = ""
            if base_yn:
                copy_string = 'EDIT_VECTOR FILENAME="POLY_COVERED<sub>*" SHAPE_TYPE=AREAS COPY_TO_NEW_LAYER=YES ' \
                              'NEW_LAYER_NAME="COVERED"\n' \
                              'EDIT_VECTOR FILENAME=POLY_COVERED<sub>* HIDDEN=YES\n'
            if nc_flag:
                copy_string = copy_string + 'EDIT_VECTOR FILENAME="POLY_NOT_COVERED<sub>*" SHAPE_TYPE=AREAS ' \
                                            'COPY_TO_NEW_LAYER=YES NEW_LAYER_NAME="NOT_COVERED"\n' \
                                            'EDIT_VECTOR FILENAME=POLY_NOT_COVERED<sub>* HIDDEN=YES\n'
            # remove imports
            for i, toe in enumerate(layer_spile_list):
                toe_info = self.toe_information[toe]
                toe_path = layer_spile_list_paths[i]
                copy_string = copy_string + 'SET_LAYER_OPTIONS FILENAME=%s HIDDEN=YES ALLOW_EXPORT=NO\n' % toe_path

            # calculate vols
            if nc_flag:
                vols_string_1 = 'EDIT_VECTOR FILENAME=NOT_COVERED SHAPE_TYPE=AREAS COPY_TO_NEW_LAYER=YES NEW_LAYER_' \
                                'NAME=ALL_AREAS\n'
                if base_yn:
                    vols_string_1 = vols_string_1 + 'EDIT_VECTOR FILENAME=COVERED SHAPE_TYPE=AREAS ' \
                                                    'COPY_TO_NEW_LAYER=YES NEW_LAYER_NAME=ALL_AREAS\n'
                vols_string = vols_string_1 + 'CALC_VOLUMES FILENAME=ALL_AREAS OUTPUT_FILENAME=%s ' \
                                              'ADD_VOLUME_ATTRS=YES\n' % output_t2t
            else:
                vols_string = 'CALC_VOLUMES FILENAME=COVERED OUTPUT_FILENAME=%s ADD_VOLUME_ATTRS=YES\n' % output_t2t
            output_t2t_shp = output_t2t.replace('.csv', '.shp')
            vols_string = vols_string + 'SET_LAYER_OPTIONS FILENAME=%s ALLOW_EXPORT=NO\n' % self.laz_cat

            vols_string = vols_string + 'EXPORT_VECTOR FILENAME=%s TYPE=SHAPEFILE GEN_PRJ_FILE=YES SHAPE_TYPE=AREAS\n' % output_t2t_shp
            if base_yn:
                # if there is in fact base coverage
                if nc_flag:
                    # if there are some not covered, hide them for now
                    vols_string = vols_string + 'SET_LAYER_OPTIONS FILENAME=NOT_COVERED HIDDEN=YES\n'
                    vols_string = vols_string + 'EDIT_VECTOR FILENAME=NOT_COVERED ATTR_TO_DELETE=*\n'
                vols_string = vols_string + 'SET_LAYER_OPTIONS FILENAME=COVERED HIDDEN=NO\n'
                vols_string = vols_string + 'EDIT_VECTOR FILENAME=COVERED ATTR_TO_DELETE=*\n'

                if nc_flag:
                    vols_string = vols_string + 'SET_LAYER_OPTIONS FILENAME=NOT_COVERED HIDDEN=YES\n'
                vols_string = vols_string + 'CALC_VOLUME_BETWEEN_SURFACES LAYER1_FILENAME=dem_files.gmc ' \
                                            'LAYER2_FILENAME=base_files.gmc AREA_FILENAME=COVERED ' \
                                            'ADD_VOLUME_ATTRS=YES OUTPUT_FILENAME=%s\n' % output_s2s
                output_s2s_shp = output_s2s.replace('.csv', '.shp')
                vols_string = vols_string + 'EXPORT_VECTOR FILENAME=%s TYPE=SHAPEFILE GEN_PRJ_FILE=YES SHAPE_TYPE=AREAS\n' % output_s2s_shp

            # screenshot
            header_string = "%s - %s" % (self.project_name, area_base)
            if not nc_flag:
                screenshot_string = 'SET_LAYER_OPTIONS FILENAME=base_files.gmc HIDDEN=YES\n'
            else:
                screenshot_string = ""
            screenshot_string = screenshot_string + 'SET_LAYER_OPTIONS FILENAME=laz_files.gmc HIDDEN=YES\n' \
                                'SET_VIEW GLOBAL_BOUNDS=%s,%s,%s,%s\n' \
                                'SET_VERT_DISP_OPTS ENABLE_HILL_SHADING=YES\n' \
                                'EXPORT_PDF FILENAME=%s DPI=100 PDF_PAGE_SIZE=A4 PDF_MARGINS=0.5,0.5,0.5,0.5 ' \
                                'PDF_HEADER="%s" PDF_FILL_PAGE=YES PDF_COMBINE_RASTERS=YES GLOBAL_BOUNDS=%s,%s,%s,%s\n' \
                                % (buff_xmin, buff_ymin, buff_xmax, buff_ymax, output_pdf_page, header_string, buff_xmin, buff_ymin, buff_xmax, buff_ymax)

            # save for later
            save_string = 'SAVE_WORKSPACE FILENAME=%s\n' \
                          'UNLOAD_ALL' % output_workspace

            # compile the strings
            process_string = start_string + import_covered_string + copy_string + vols_string + screenshot_string + save_string

            if not os.path.isdir(self.csv_folder):
                os.mkdir(self.csv_folder)

            with open(gms_name, 'w') as g:
                g.write(process_string)

            subprocess.call([self.gm_exe, gms_name])

    def import_toe_shp(self, test_path):
        probe = shp.Reader(test_path)
        dict_of_dicts = {}
        for i, record in enumerate(probe.records()):
            name = record.as_dict()['NAME']
            dict_of_dicts[name] = record.as_dict()

        shp_panda = pd.DataFrame.from_dict(dict_of_dicts, orient='index')
        try:
            shp_panda_filtered = shp_panda[["NAME", 'NET_VOLUME', 'CUT_VOLUME', 'CUT_AREA',
                                            'FILL_VOLUM', 'FILL_AREA', 'NAME']]
        except KeyError:
            print('bugger')
            print(shp_panda.columns)
            print(test_path)
            exit()
        try:
            shp_panda_filtered.loc[:, 'CUT_AREA_M'] = shp_panda_filtered['CUT_AREA'].apply(lambda x: float(x) * 1)
            # shp_panda_filtered.loc[:, 'CUT_AREA_M'] = shp_panda_filtered['CUT_AREA'].apply(lambda x: float(x) * 1000000)
            shp_panda_filtered.loc[:, 'FILL_AREA_M'] = shp_panda_filtered['FILL_AREA'].apply(lambda x: float(x) * 1)
            # shp_panda_filtered.loc[:, 'FILL_AREA_M'] = shp_panda_filtered['FILL_AREA'].apply(lambda x: float(x) * 1000000)
            shp_panda_filtered.loc[:, 'CUT_VOLUME'] = shp_panda_filtered['CUT_VOLUME'].astype(float)
            shp_panda_filtered.loc[:, 'AV_THICKNESS'] = shp_panda_filtered['CUT_VOLUME'] / shp_panda_filtered['CUT_AREA_M']
            shp_panda_filtered.drop(columns=['CUT_AREA', 'FILL_AREA'], inplace=True)
        except KeyError:
            shp_panda_filtered = None
            print(shp_panda.columns)
            print('issue')
            print(test_path)
            exit()
        return shp_panda_filtered

    def csv_tool(self):
        # no longer uses csv due to encoding issue!
        print('Compiling volumes...')
        for area_name in self.csv_dic.keys():
            area_info = self.csv_dic[area_name]
            s2s_orginal_csv = area_info['S2S']
            t2t_orginal_csv = area_info['T2T']
            s2s_shp = s2s_orginal_csv.replace('.csv', '.shp')
            t2t_shp = t2t_orginal_csv.replace('.csv', '.shp')
            s2s_panda = None
            t2t_panda = None
            s2s_yes = False
            t2t_yes = False
            if os.path.isfile(s2s_shp):
                s2s_panda = self.import_toe_shp(s2s_shp).add_suffix('_S2S')
                s2s_yes = True
            if os.path.isfile(t2t_shp):  # should bloody well be
                t2t_panda = self.import_toe_shp(t2t_shp).add_suffix('_T2T')
                t2t_yes = True

            # merge the two pandas
            merged_success = True
            just_one = False
            if s2s_yes and t2t_yes:
                print('merging T2T and S2S results')
                try:
                    merged_results = t2t_panda.join(s2s_panda, sort=False, how='left')
                except Exception as e:
                    print('Error merging tables, ', e)
                    merged_results = None
                    merged_success = False
            else:
                just_one = True
                print('Only one calculation style found')
                merged_results = t2t_panda
                merged_success = True


            # add calculation method
            method_list = []
            if not just_one:
                try:
                    for item in merged_results['CUT_VOLUME_S2S'].to_list():
                        if math.isnan(item):
                            method_list.append('Toe')
                        else:
                            method_list.append('Base')
                except KeyError:
                    method_list = ['Toe' for item in merged_results['Cut_Volume_T2T'].to_list()]
            else:
                method_list = ['Toe' for item in merged_results['CUT_VOLUME_T2T'].to_list()]
            merged_results['method'] = method_list

            # get times
            temp = merged_results
            merged_results = temp.join(self.toes_times_pd.T, sort=False, how='left')

            # export this to a csv for future records
            interim_csv_name = "%s_INTERIM.csv" % area_name
            interim_csv_path = os.path.join(self.csv_folder, interim_csv_name)
            if merged_success:
                merged_results.to_csv(interim_csv_path)

            # create "clean" table where s2s has overwritten t2t where possible.
            if not just_one:
                try:
                    merged_results['CUT_VOLUME_S2S'].fillna(merged_results['CUT_VOLUME_T2T'], inplace=True)
                    merged_results['CUT_AREA_M_S2S'].fillna(merged_results['CUT_AREA_M_T2T'], inplace=True)
                    merged_results['FILL_AREA_M_S2S'].fillna(merged_results['FILL_AREA_M_T2T'], inplace=True)
                    merged_results['FILL_VOLUM_S2S'].fillna(merged_results['FILL_VOLUM_T2T'], inplace=True)
                    merged_results['NET_VOLUME_S2S'].fillna(merged_results['NET_VOLUME_T2T'], inplace=True)
                    merged_results['AV_THICKNESS_S2S'].fillna(merged_results['AV_THICKNESS_T2T'], inplace=True)
                except KeyError:
                    merged_success = False
                    print('Issue on ', area_name)
                    print(merged_results.columns)
                    exit()

            # rename columns appropriately, drop unnecessary
            if merged_success:
                issue = False
                if not just_one:
                    # print(merged_results.columns)
                    try:
                        print(merged_results.columns)
                        if 'method' in merged_results.columns:
                            merged_results = merged_results[['human_time', 'method', 'NET_VOLUME_S2S', 'CUT_VOLUME_S2S', 'CUT_AREA_M_S2S', 'FILL_VOLUM_S2S', 'FILL_AREA_M_S2S', 'AV_THICKNESS_S2S']]
                        else:
                            merged_results = merged_results[['human_time', 'Method', 'NET_VOLUME_S2S', 'CUT_VOLUME_S2S', 'CUT_AREA_M_S2S', 'FILL_VOLUM_S2S', 'FILL_AREA_M_S2S', 'AV_THICKNESS_S2S']]
                    except TypeError:
                        col_list = ['NAME_S2S', 'human_time', 'method', 'NET_VOLUME_S2S', 'CUT_VOLUME_S2S', 'CUT_AREA_M_S2S', 'FILL_VOLUM_S2S', 'FILL_AREA_M_S2S', 'AV_THICKNESS_S2S']
                        for c in col_list:
                            if c not in merged_results.columns:
                                print('Column %s is culprit' % c)
                        merged_results = None
                        issue = True
                else:
                    try:
                        merged_results = merged_results[['human_time', 'method','NET_VOLUME_T2T', 'CUT_VOLUME_T2T', 'CUT_AREA_M_T2T', 'FILL_VOLUM_T2T', 'FILL_AREA_M_T2T', 'AV_THICKNESS_T2T']]
                    except TypeError:
                        merged_results = None
                        issue = True
                print('Just one?', just_one)
                if not issue:
                    if not just_one:
                        merged_results = merged_results.rename(columns={'CUT_VOLUME_S2S': 'Cut volume, m3',
                                                       'CUT_AREA_M_S2S': 'Cut area, m2',
                                                       'FILL_AREA_M_S2S': 'Fill area, m2',
                                                       'FILL_VOLUM_S2S': 'Fill volume, m3',
                                                       'NAME_S2S': 'Stockpile',
                                                       'NET_VOLUME_S2S': 'Net volume, m3',
                                                       'AV_THICKNESS_S2S': 'Av thickness, m',
                                                       'human_time': 'Survey time (GMT)',
                                                       'method': 'Method'})
                    else:
                        merged_results = merged_results.rename(columns={'CUT_VOLUME_T2T': 'Cut volume, m3',
                                                       'CUT_AREA_M_T2T': 'Cut area, m2',
                                                       'FILL_AREA_M_T2T': 'Fill area, m2',
                                                       'FILL_VOLUM_T2T': 'Fill volume, m3',
                                                       'NAME_T2T': 'Stockpile',
                                                       'NET_VOLUME_T2T': 'Net volume, m3',
                                                       'AV_THICKNESS_T2T': 'Av thickness, m',
                                                       'human_time': 'Survey time (GMT)'})

                    # export merged csv
                    end_filepath = os.path.join(self.csv_folder, "%s_FINAL.csv" % area_name)
                    merged_results.to_csv(end_filepath)
                else:
                    print('No merged_results for ', area_name)


# option to rerun from CSV stage, single csv!
class report_make:
    def __init__(self, noot):
        print('Step 3 - Reporting - selected')
        self.csv_folder = noot['csv folder']
        self.project_name = noot['project_name']
        self.client_name = noot['client_name']
        self.client_company = noot['client_company']
        self.user_name = noot['user_name']
        self.product_type = noot['product_type']
        self.show_all_vals = noot['show_all_vals']

        self.output_folder = os.path.join(self.csv_folder, "Output")
        if not os.path.isdir(self.output_folder):
            os.mkdir(self.output_folder)

        self.final_csv_list = [os.path.join(self.csv_folder, a) for a in os.listdir(self.csv_folder)
                               if a.endswith('_FINAL.csv')]
        self.merged_csv_list = [os.path.join(self.csv_folder, a) for a in os.listdir(self.csv_folder)
                                if a.endswith('_MERGED.csv')]
        self.pdf_pages_list = [os.path.join(self.csv_folder, a) for a in os.listdir(self.csv_folder)
                               if a.endswith('.pdf')]

        # technically only the final matter most of the time
        self.final_pandas_dic = {}
        if not self.show_all_vals:
            for csv_file in self.final_csv_list:
                area_base_name = os.path.splitext(os.path.basename(csv_file))[0]
                area_panda = pd.read_csv(csv_file).round(decimals=2)
                # enforce column order
                if 'Cut_Volume_m3' not in area_panda.columns.tolist():
                    area_panda.reset_index()
                    self.final_pandas_dic[area_base_name] = area_panda
                else:
                    print(area_panda.columns)
                    if 'Stockpile Name' not in area_panda.columns.to_list():
                        area_panda.rename(columns={'Unnamed: 0': 'Stockpile Name'}, inplace=True)
                    area_panda.rename(columns={'index': 'Stockpile Name'}, inplace=True)
                    try:
                        area_panda = area_panda[['Stockpile Name', 'Time_GMT', 'Method', 'Net_Volume_m3', 'Cut_Volume_m3',
                                             'Cut_Area_m2', 'Fill_Area_m2', 'Fill_Volume_m3', 'Total_Volume_m3',
                                             'Average_thickness_m']]
                    except KeyError:
                        print(area_panda.columns)
                        exit()
                    area_panda_2 = area_panda.rename(columns={'Cut_Volume_m3': 'Cut Volume (m3)',
                                                              'Time_GMT': 'Time (GMT)',
                                                              'Cut_Area_m2': 'Cut Area (m2)',
                                                              'Fill_Area_m2': 'Fill Area (m2)',
                                                              'Fill_Volume_m3': 'Fill Volume (m3)',
                                                              'Total_Volume_m3': 'Total Volume (m3)',
                                                              'Average_thickness_m': 'Average thickness (m)',
                                                              'Net_Volume_m3': 'Net Volume (m3)'})
                    self.final_pandas_dic[area_base_name] = area_panda_2
        else:
            for csv_file in self.merged_csv_list:
                area_base_name = os.path.splitext(os.path.basename(csv_file))[0]
                area_panda = pd.read_csv(csv_file).round(decimals=2)
                print(area_base_name, area_panda.columns)
                print('############')
                print('Show all vals not yet supported')




            pass

        self.word_doc = os.path.join(self.output_folder, "%s_Report.docx" % self.project_name)

        # constructs word doc (?)
        # https://stackoverflow.com/questions/40596518/writing-a-python-pandas-dataframe-to-word-document
        for item in self.pdf_pages_list:
            base_pdf = os.path.basename(item)
            new_path = os.path.join(self.output_folder, base_pdf)
            shutil.copy2(item, self.output_folder)
        self.make_word_doc()

    def make_word_doc(self):
        # variables to fill:
        # {{date_input}}, {{client_fullname}}, {{client_company}}, {{client_firstname}}, {{product}}, {{tables}},
        # {{user_name}}
        client_first_name = self.client_name.split()[0]
        doc = docxtpl.DocxTemplate(template_docx)
        styles = doc.styles
        for s in styles:
            print('Style', s.name, s.type)
        doco = doc.new_subdoc()
        date_now = datetime.datetime.today().strftime('%c')
        tables_dic = {}
        tables_headers = {}
        for table_setup in self.final_pandas_dic.keys():
            title = table_setup.replace('_', ' ').replace('FINAL', '').replace('Toes', '')
            table_pd = self.final_pandas_dic[table_setup]
            table_pd.drop_duplicates(subset=table_pd.columns.to_list(), keep='first', inplace=True)
            doco.add_paragraph("")
            doco.add_heading(title, level=1)
            tin = doco.add_table(table_pd.shape[0]+1, table_pd.shape[1])
            tin.allow_autofit = True
            for j in range(table_pd.shape[-1]):
                tin.cell(0,j).text = table_pd.columns[j]
            for i in range(table_pd.shape[0]):
                for j in range(table_pd.shape[-1]):
                    tin.cell(i+1,j).text = str(table_pd.values[i,j])
            for row in tin.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(9)

        context = {'date_input': date_now, 'client_fullname': self.client_name, 'client_company': self.client_company,
                   'client_firstname': client_first_name, 'product': self.product_type.lower().capitalize(),
                   'tables': doco, 'user_name': self.user_name}
        doc.render(context)
        doc.save(self.word_doc)
        print("I have theoretically saved to", self.word_doc)

        pass






# special VOXEL MODE running purely on lastools.
class voxel_mode:
    def __init__(self, noot):
        print('Ninja Mode selected. Good luck!')
        pass

    def set_up(self):
        # creates geojson of bb shp
        # lasclip to
        pass


class pingu:
    def __init__(self):
        noot = goo()
        print(noot.__dict__)
        # detect what stage it is at:
        if 'Step_1' in noot.command:
            prepare(noot)
        elif 'Step_2' in noot.command:
            calculate(noot)
        elif 'Step_3' in noot.command:
            report_dic = {'client_name': noot.client_name, 'client_company': noot.client_company,
                          'user_name': noot.user_name, 'csv folder': noot.csv_folder, 'project_name': noot.site_name,
                          'product_type': noot.product_type, 'show_all_vals': noot.show_all_vals}
            report_make(report_dic)
        elif 'Ninja' in noot.command:
            voxel_mode(noot)
        else:
            print('Error in interpretation of step')
            exit()


@Gooey(program_name="Insert PSID - v2", use_legacy_titles=True, required_cols=1, default_size=(1050, 700))
def goo():
    parser = GooeyParser(description="Simple volumes worker")
    sub_pars = parser.add_subparsers(help='commands', dest='command')

    # -- First step: gets volumes areas prepped -- #
    setup_step = sub_pars.add_parser('Step_1', help='Run volumes preparation')
    # needs las files, plus shp showing roughly where the spiles are. it will buff these.
    setup_step.add_argument('-las_folder', metavar="LAS Folder", widget='DirChooser',
                            help='Select folder containing classified and adjusted las files')
    setup_step.add_argument('-las_json', metavar="LAS Tile Layout", widget='FileChooser',
                            help='Select tile layout json for the las files')
    setup_step.add_argument('-pre_shp_folder', metavar="Directory with SHP in rough location of spiles",
                            widget="DirChooser", help="If no guide polygons are given, entire site will be prepped.")  # if nothing here, grids whole site.
    setup_step.add_argument('-buffer_size', metavar="Amount by which to buffer these polygons (m)", default=500)
    setup_step.add_argument('-cores', metavar="Cores to use for gridding and contours", default=19)
    setup_step.add_argument('-tile_size', metavar="Tile size (m)", default=500)
    setup_step.add_argument('-proj', metavar="Project projection system",
                            choices=['MGA_GDA94', 'MGA_GDA2020', 'AMG_AGD84', 'AMG_AGD66', 'ELL'], default='MGA_GDA94')
    setup_step.add_argument('-zone', metavar="Project zone",
                            choices=zones, default='56')
    setup_step.add_argument('-open_gmws', metavar="Open GMW on completion?", action="store_true")
    setup_step.add_argument('-gm_exe', metavar="Global Mapper location", widget='FileChooser',
                            help='Select location for your Global Mapper .exe file',
                            default=r"C:\Program Files\GlobalMapper21.1_64bit\global_mapper.exe")
    setup_step.add_argument('-quick_name', metavar="Please add a quick name for your area",
                            help="This is just for aesthetics and will not be used in final report.")

    # -- Second step: Calculate volumes from saved GMW -- #
    calculate_step = sub_pars.add_parser('Step_2', help="Run calculations")
    # needs working GMW file, gm exe, optional base lasfolder
    # optional shapefile containing areas into which to divide pandas (else will divide by base)
    calculate_step.add_argument('-gmw_file', metavar="GMW Working File", widget='FileChooser',
                                help='Select the working GMW file')
    calculate_step.add_argument('--name', metavar='FORMAL name for project')
    calculate_step.add_argument('-gm_exe', metavar="Global Mapper location", widget='FileChooser',
                                help='Select location for your Global Mapper .exe file',
                                default=r"C:\Program Files\GlobalMapper21.1_64bit\global_mapper.exe")
    calculate_step.add_argument('-base_folder', metavar="Select folder with base files.",
                                help="This can be left blank. They must be in GMG format.", widget="DirChooser")
    calculate_step.add_argument('-laz_folder', metavar="Select folder with original laz files.", widget="DirChooser")
    calculate_step.add_argument('-dem_folder', metavar="Select folder with dem files calculated last step.", widget="DirChooser")

    calculate_step.add_argument('-cores', metavar="Cores to use for gridding and contours", default=19)
    calculate_step.add_argument('-include_step_3',
                                metavar="Do you want to pre-emptively create reports?",
                                help="i.e. do you trust this site to behave in this script?", action='store_true')
    calculate_step.add_argument('-toe_technique',
                                metavar="Tick if your toes had elevations set by the client.", action='store_true')
    calculate_step.add_argument('-proj', metavar="Project projection system",
                                choices=['MGA_GDA94', 'MGA_GDA2020', 'AMG_AGD84', 'AMG_AGD66', 'ELL'], default='MGA_GDA94')
    calculate_step.add_argument('-zone', metavar="Project zone",
                                choices=zones, default='56')
    calculate_step.add_argument('-client_name', metavar="Who should this be addressed to?")
    calculate_step.add_argument('-client_company', metavar="For whom do they work?")
    calculate_step.add_argument('-user_name', metavar="What is your name?")
    calculate_step.add_argument('-product_type', metavar="What's in those spiles?")
    calculate_step.add_argument('-show_all_vals', metavar="Do you want both T2T and S2S values shown in the report?",
                                action="store_true")

    # -- Third step: Build reports from CSVs and GMW -- #
    # looks in folder for files matching patterns of auto exports
    report_step = sub_pars.add_parser('Step_3', help="Create reports")
    report_step.add_argument('-csv_folder', metavar="Select the folder containing your volume csv files.",
                             help="Please ensure naming follows guidelines.", widget='DirChooser')
    # will insert table with both t2t and s2s if both are available
    report_step.add_argument('-show_all_vals', metavar="Do you want both T2T and S2S values shown in the report?",
                             action="store_true")
    report_step.add_argument('-site_name', metavar="Please enter the site name")
    '''report_step.add_argument('-month', metavar="Please enter the date code")'''  # tf is this
    report_step.add_argument('-client_name', metavar="What is the client name?", help="e.g. 'Bob Ross'")
    report_step.add_argument('-client_company', metavar='For whom do they work?')
    report_step.add_argument('-user_name', metavar="What is your name?", help="e.g. 'Jasnah Kholin'")
    report_step.add_argument('-product_type', metavar="What's in those spiles?")

    # -- NINJA MODE -- #
    # Special voxel-based technique (potential "next workflow") to be a) tested and b) used on sites with
    # pre existing toes
    ninja_mode_option = sub_pars.add_parser('Ninja_Mode', help='Special lastools workflow STILL IN TESTING')
    ninja_mode_option.add_argument('-input_las', metavar="Folder of input las", widget='DirChooser')
    ninja_mode_option.add_argument('-input_json', metavar="Input las geojson tile layout", widget='FileChooser')
    ninja_mode_option.add_argument('-base_las', metavar="Folder of base las", widget='DirChooser')
    ninja_mode_option.add_argument('-toe_shp_folder', metavar="Folder of toe shapefiles", widget='DirChooser')
    ninja_mode_option.add_argument('-bb_shp_folder', metavar="Folder of toe areas", widget='DirChooser')
    ninja_mode_option.add_argument('-last_mnth', metavar="Folder of last month's data", widget='DirChooser')
    ninja_mode_option.add_argument('-site_name', metavar="Name of site", help="Will go on final report")
    ninja_mode_option.add_argument('-client_name', metavar="Name of client",
                                   help="e.g. Bob Smith. Will go on final report")
    ninja_mode_option.add_argument('-user_name', metavar="Your name", help="e.g. Bob Smith. Will go on final report")

    return parser.parse_args()


if __name__ == "__main__":
    pingu()





